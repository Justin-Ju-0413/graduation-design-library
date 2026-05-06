"""Extract text from PDF and DOCX files in the FYP requirement directory."""

import os, sys

BASE = r"C:\Users\16084\Documents\Graduation_Design_Library\11_FYP_requirement"

# ── 1. Guideline PDF ──────────────────────────────────────────────────────────
guideline_pdf = os.path.join(BASE, "3+1+X FYP_Guideline_251024_final (2).pdf")
guideline_out = os.path.join(BASE, "guideline_extracted.txt")

print("=" * 70)
print("Extracting: Guideline PDF")
print("=" * 70)

try:
    import fitz  # PyMuPDF
    doc = fitz.open(guideline_pdf)
    print(f"  Pages: {doc.page_count}")
    lines = []
    for i, page in enumerate(doc):
        text = page.get_text()
        lines.append(f"--- Page {i+1} ---\n{text}")
    full = "\n".join(lines)
    with open(guideline_out, "w", encoding="utf-8") as f:
        f.write(full)
    print(f"  Written: {guideline_out}  ({len(full)} chars)")
    doc.close()
except Exception as e:
    print(f"  PyMuPDF failed: {e}")
    # fallback to pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(guideline_pdf) as pdf:
            print(f"  Pages (pdfplumber): {len(pdf.pages)}")
            lines = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                lines.append(f"--- Page {i+1} ---\n{text}")
            full = "\n".join(lines)
            with open(guideline_out, "w", encoding="utf-8") as f:
                f.write(full)
            print(f"  Written: {guideline_out}  ({len(full)} chars)")
    except Exception as e2:
        print(f"  pdfplumber also failed: {e2}")

# ── 2. Timeline PDF ───────────────────────────────────────────────────────────
timeline_pdf = os.path.join(BASE, "AY202526_3+1+X_FYP_Timeline_251024.pdf")
timeline_out = os.path.join(BASE, "timeline_extracted.txt")

print("\n" + "=" * 70)
print("Extracting: Timeline PDF")
print("=" * 70)

try:
    import fitz
    doc = fitz.open(timeline_pdf)
    print(f"  Pages: {doc.page_count}")
    lines = []
    for i, page in enumerate(doc):
        text = page.get_text()
        lines.append(f"--- Page {i+1} ---\n{text}")
    full = "\n".join(lines)
    with open(timeline_out, "w", encoding="utf-8") as f:
        f.write(full)
    print(f"  Written: {timeline_out}  ({len(full)} chars)")
    doc.close()
except Exception as e:
    print(f"  PyMuPDF failed: {e}")
    try:
        import pdfplumber
        with pdfplumber.open(timeline_pdf) as pdf:
            lines = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                lines.append(f"--- Page {i+1} ---\n{text}")
            full = "\n".join(lines)
            with open(timeline_out, "w", encoding="utf-8") as f:
                f.write(full)
            print(f"  Written: {timeline_out}  ({len(full)} chars)")
    except Exception as e2:
        print(f"  pdfplumber also failed: {e2}")

# ── 3. DOCX Template ──────────────────────────────────────────────────────────
template_docx = os.path.join(BASE, "3+1+X FYP_Report_Template_251024.docx")
template_out = os.path.join(BASE, "template_extracted.txt")

print("\n" + "=" * 70)
print("Extracting: DOCX Template")
print("=" * 70)

try:
    from docx import Document
    doc = Document(template_docx)
    lines = []
    # Paragraphs
    for p in doc.paragraphs:
        lines.append(p.text)
    # Tables
    for ti, table in enumerate(doc.tables):
        lines.append(f"\n--- Table {ti+1} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    full = "\n".join(lines)
    with open(template_out, "w", encoding="utf-8") as f:
        f.write(full)
    print(f"  Written: {template_out}  ({len(full)} chars)")
except Exception as e:
    print(f"  python-docx failed: {e}")
    # fallback: unzip and read XML
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(template_docx, "r") as z:
            xml_content = z.read("word/document.xml")
        root = ET.fromstring(xml_content)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        texts = []
        for t in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
            if t.text:
                texts.append(t.text)
        full = "\n".join(texts)
        with open(template_out, "w", encoding="utf-8") as f:
            f.write(full)
        print(f"  Written (XML fallback): {template_out}  ({len(full)} chars)")
    except Exception as e2:
        print(f"  XML fallback also failed: {e2}")

print("\n" + "=" * 70)
print("DONE")
print("=" * 70)
