"""Fill FYP template cover pages (first 3 pages) with student info.
Output a clean DOCX with: cover page, declaration, abstract + TOC only."""
from pathlib import Path
from docx import Document
from lxml import etree

TEMPLATE = Path(r"C:\Users\16084\Documents\Graduation_Design_Library\11_FYP_requirement\3+1+X FYP_Report_Template_251024.docx")
OUT_DOCX = Path(r"C:\Users\16084\Documents\Graduation_Design_Library\11_FYP_requirement\FYP_COVER.docx")

TITLE = "RISC-V Custom Instruction Based Lightweight CNN Accelerator FPGA Prototype Validation"
STUDENT_NAME = "JU JIAXING"
HOME_UNIV = "South China University of Technology"
STUDENT_ID = "2501026"

doc = Document(str(TEMPLATE))

# =====================================================
# PASS 1: Fill in user info on cover page
# =====================================================
for p in doc.paragraphs:
    text = p.text.strip()

    # Proposal Code: remove "Proposal Code" placeholder
    if "2025/26" in p.text and "Proposal Code" in p.text:
        for run in p.runs:
            run.text = run.text.replace("-Proposal Code", "").replace("Proposal Code", "")
        # Clean trailing "-"
        for run in p.runs:
            run.text = run.text.rstrip("-").rstrip()
        # If all runs empty, set first run to "2025/26"
        if all(not r.text.strip() for r in p.runs):
            p.runs[0].text = "2025/26"

    # Project Title placeholder (yellow highlighted)
    if text == "Project Title":
        for run in p.runs:
            run.font.highlight_color = None
            run.text = ""
        if p.runs:
            p.runs[0].text = TITLE

    # Student Name:
    if text == "Student Name:":
        for run in p.runs:
            if "Student Name:" in run.text:
                run.text = f"Student Name: {STUDENT_NAME}"
                break

    # Home University:
    if text == "Home University:":
        for run in p.runs:
            if "Home University:" in run.text:
                run.text = f"Home University: {HOME_UNIV}"
                break

    # Student ID:
    if text == "Student ID:":
        for run in p.runs:
            if "Student ID:" in run.text:
                run.text = f"Student ID: {STUDENT_ID}"
                break

# =====================================================
# PASS 2: Fill declaration table
# =====================================================
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                t = p.text.strip()
                if t == "Project Title:":
                    for run in p.runs:
                        if "Project Title:" in run.text:
                            run.text = f"Project Title: {TITLE}"
                            break
                elif t == "Student Name:":
                    for run in p.runs:
                        if "Student Name:" in run.text:
                            run.text = f"Student Name: {STUDENT_NAME}"
                            break
                elif t == "Student ID:":
                    for run in p.runs:
                        if "Student ID:" in run.text:
                            run.text = f"Student ID: {STUDENT_ID}"
                            break

# =====================================================
# PASS 3: XML-level cleanup — remove instruction paragraphs
# and everything from "1. Introduction" onward
# =====================================================
NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
body = doc.element.body

# Tags that indicate removable content
REMOVE_KEYWORDS = [
    "(Parts in yellow highlights should be updated/replaced and unhighlighted)",
    "(for reference, please remove it from your report)",
    "(Please note that the Background",
    "(Please note that the term",
    "(Please note that this",
    "guiding checklist",
    "The Introduction introduces the topic",
    "A methodology describes the research",
    "This section provides the full context",
    "The results section summarizes",
    "The discussion section analyses",
    "The conclusion should answer",
    "(Please remove this part",
    "(or more if needed)",
    "The following shows the format",
    "Make sure that you properly reference",
    "(Taken from:",
    "www.ieeetclt.org",
    "Basic format for",
    "*Basic format",
    # Remove all TOC entries (keep Abstract heading)
    "Abstract i", "Acknowledgements ii", "Contents iii",
    "List of Figures iv", "List of Tables v",
    "Acknowledgements Contents",
    "Figure 1", "Figure 2", "Figure 3",
    "Table 1", "Table 2", "Table 3",
    "e.g. P.", "MIMO channel", "General Frequency",
    "Block Diagram", "Complexity of the",
    "List of Figures",
    "List of Tables",
    "Have you:",
    # Remove TOC tab-leader entries (Abstract .... i etc.)
    "Abstract\t", "Acknowledgements\t", "Contents\t",
    "List of Figures\t", "List of Tables\t",
    "briefly stated the background",
    "told what the problem is",
    "explained why it was necessary",
    "described how the study was conducted",
    "outlined the main results",
    "given a brief conclusion",
    "An abstract is an overview",
    "introduced the topic or the problem",
    "identified the limitations of previous",
    "stated the objectives of the study",
    "outlined the methods of the study",
    "established the importance of the problem",
    "identified knowledge gaps",
    "justified the purpose",
    "reviewed the literature",
    "described how the problem was studied",
    "described how the research was performed",
    "explained how the data was analyzed",
    "used past tense to describe",
    "used present tense to refer",
    "labelled all graphs",
    "introduced and correctly labelled",
    "summarized key findings in the graphs",
    "analyzed and interpreted the key findings",
    "correctly labelled any graphs",
    "stated how successfully objectives were met",
    "stated the implications of the research",
    "suggested any future research",
]

# Find all paragraphs, mark those to remove
to_remove = []
intro_found = False

for p in doc.paragraphs:
    text = p.text.strip()

    # Flag: "1. Introduction" → remove this and everything after
    if text == "1. Introduction" or text == "1.Introduction" or text.startswith("1. Introduction"):
        intro_found = True

    if intro_found:
        to_remove.append(p._element)
        continue

    # Check all remove keywords
    for kw in REMOVE_KEYWORDS:
        if kw.lower() in text.lower():
            to_remove.append(p._element)
            break

    # Remove abstract checklist numbered items (1-6 in parentheses)
    if text in [f"{i})" for i in range(1, 8)]:
        to_remove.append(p._element)

    # Remove empty paragraphs that follow removed checklist items
    # (handled by the keyword matching above)

# Remove in reverse document order
for elem in to_remove:
    try:
        body.remove(elem)
    except Exception:
        pass

# =====================================================
# PASS 4: Remove List of Figures / List of Tables example rows
# and clean up TOC area
# =====================================================
# Remove example table rows (Figure 1, Figure 2, Table 1, Table 2, Figure 3..., Table 3..., etc.)
# Also remove entire tables that contain only placeholder content
tables_to_remove = []
for table in doc.tables:
    rows_to_remove = []
    for i, row in enumerate(table.rows):
        cell_text = " ".join(c.text for c in row.cells)
        if any(kw in cell_text for kw in ["e.g. P.", "Figure 1", "Figure 2", "Table 1", "Table 2",
                                            "MIMO channel", "General Frequency",
                                            "Block Diagram of Frequency",
                                            "Complexity of the Exhaustive",
                                            "Title                                         e.g.",
                                            "[1] ", "[2] ", "[3] ", "[4] ", "[5] ",
                                            "[6] ", "[7] ", "[8] ", "[9] ", "[10] ",
                                            "[11] ", "[12] ", "[13] ", "[14] ", "[15] ",
                                            "[16] ", "[17] ", "[18] ", "[19] ", "[20] ",
                                            "[21] ", "[22] ", "[23] ", "[24] ", "[25] ",
                                            "J. K. Author",
                                            "G. O. Young",
                                            "W.-K. Chen",
                                            "J. U. Duncombe",
                                            "E. P. Wigner",
                                            "E. H. Miller",
                                            "E. E. Reber",
                                            "J. H. Davis",
                                            "J. Jones",
                                            "R. J. Vidmar",
                                            "S. L. Talleen",
                                            "A. Harriman",
                                            "D. B. Payne",
                                            "D. Ebehard",
                                            "G. Brandli",
                                            "J. O. Williams",
                                            "N. Kawasaki",
                                            "A. Harrison",
                                            "B. Smith",
                                            "A. Brahms",
                                            "PROCESS Corp",
                                            "Musical toothbrush",
                                            "*Transmission Systems*",
                                            "*Motorola Semiconductor*",
                                            "IEEE Criteria",
                                            "Letter Symbols",
                                            "*Name of Manual*",
                                            "City of Publisher",
                                            "*Plastics*",
                                            "Available: http",
                                            "U.S. Patent",
                                            "Unpublished",
                                            "to be published",
                                            "private communication",
                                            "Belmont, CA",
                                            "Los Angeles, CA",
                                            "Austin,",
                                            "Winston-Salem",
                                            "Phoenix, AZ",
                                            "Cambridge, MA",
                                            "Osaka, Japan",
                                            "Stuttgart, Germany",
                                            "Humanist",
                                            "NEXIS Library",
                                            "Computer Group Repository",
                                            ")]{.underline}",
                                            ")]{.mark}",
                                            ]):
            rows_to_remove.append(row._element)

    for elem in rows_to_remove:
        try:
            table._element.remove(elem)
        except Exception:
            pass

# Remove tables that only contain "Figure 3..." or "Table 3..." placeholder content
for table in doc.tables:
    all_cell_text = ""
    for row in table.rows:
        for cell in row.cells:
            all_cell_text += cell.text
    if "Figure 3" in all_cell_text or "Table 3" in all_cell_text:
        try:
            body.remove(table._element)
        except Exception:
            pass

# =====================================================
# PASS 4: Remove excessive empty paragraphs between declaration and copyright
# =====================================================
# The template has ~14 empty paragraphs between the declaration text and the
# copyright notice. Keep only a few for spacing.
empty_between_decl_copyright = []
found_decl = False
found_copyright = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "I have read the student handbook" in text:
        found_decl = True
        continue
    if "No part of this report may be reproduced" in text:
        found_copyright = True
        continue
    if found_decl and not found_copyright and not text:
        empty_between_decl_copyright.append(p._element)

# Keep first 2 empty paragraphs, remove the rest
for elem in empty_between_decl_copyright[2:]:
    try:
        body.remove(elem)
    except Exception:
        pass

# Also remove large-space paragraphs after copyright (space_after=127000)
for p in doc.paragraphs:
    if p.paragraph_format.space_after and p.paragraph_format.space_after >= 100000:
        try:
            body.remove(p._element)
        except Exception:
            pass

# Remove remaining TOC entries (style "toc 1")
for p in doc.paragraphs:
    if p.style and p.style.name and "toc" in p.style.name.lower():
        try:
            body.remove(p._element)
        except Exception:
            pass

# Also remove remaining empty paragraphs with large spacing
for p in doc.paragraphs:
    text = p.text.strip()
    if not text and p.paragraph_format.space_after and p.paragraph_format.space_after >= 50000:
        try:
            body.remove(p._element)
        except Exception:
            pass

# =====================================================
# Save
# =====================================================
doc.save(str(OUT_DOCX))
print(f"Saved: {OUT_DOCX}")
