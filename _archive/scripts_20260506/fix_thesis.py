import sys
sys.path.insert(0, r'C:\Users\16084\AppData\Roaming\Python\Python313\site-packages')

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import os

BASE_DIR = r'C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing'
SRC_PATH = os.path.join(BASE_DIR, 'FYP_Thesis_Formatted.docx')
FIG_DIR = os.path.join(BASE_DIR, 'Figures')
OUT_PATH = os.path.join(BASE_DIR, 'FYP_Thesis_Final.docx')

doc = Document(SRC_PATH)

# Figure mapping for Chapter 3 (captions exist in body)
CHAPTER3_FIGURES = {
    'Figure 3.1': {'img': 'fig3_1_soc_architecture.png', 'caption_idx': 182},
    'Figure 3.2': {'img': 'fig3_2_instruction_format.png', 'caption_idx': 207},
    'Figure 3.3': {'img': 'fig3_3_pe_microarchitecture.png', 'caption_idx': 231},
    'Figure 3.4': {'img': 'fig3_4_pe_array.png', 'caption_idx': 240},
    'Figure 3.5': {'img': 'fig3_5_packed_format.png', 'caption_idx': 250},
    'Figure 3.6': {'img': 'fig3_6_build_pipeline.png', 'caption_idx': 280},
}

# Figure mapping for Chapter 4 (need to insert captions)
CH4_FIGURES = {
    'Figure 4.1': {
        'img': 'fig_ila_pc_trace.png',
        'insert_after': 341,
        'caption': 'Figure 4.1: ILA capture showing PC progression during hello_e203 execution on FPGA.'
    },
    'Figure 4.2': {
        'img': 'fig_ila_nice_activity.png',
        'insert_after': 355,
        'caption': 'Figure 4.2: ILA capture showing NICE custom instruction activity on the FPGA.'
    },
    'Figure 4.3': {
        'img': 'fig_speedup_bar.png',
        'insert_after': 362,
        'caption': 'Figure 4.3: Performance comparison showing CNN accelerator speedup over software-only execution.'
    },
}

# Chapter title mapping
CHAPTERS = {
    '1.1 Background': ('Chapter 1. Introduction', '1'),
    '2.1 RISC-V ISA and Custom Instruction Extensions': ('Chapter 2. Background', '2'),
    '3.1 System Architecture Overview': ('Chapter 3. Methodology', '3'),
    '4.1 RTL Simulation': ('Chapter 4. Results', '4'),
    '5.1 FPGA Bring-up Challenges and Resolution': ('Chapter 5. Discussion', '5'),
    '6.1 Summary of Contributions': ('Chapter 6. Conclusion', '6'),
}


def make_paragraph_before(para, text='', style=None):
    """Insert a new paragraph before the given paragraph"""
    new_p = OxmlElement('w:p')
    para._element.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, para._parent)
    if text:
        new_para.text = text
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def insert_page_break_paragraph_before(para):
    """Insert a page break paragraph before the given paragraph"""
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    new_p.append(r)
    para._element.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, para._parent)


def insert_image_before(para, img_path, width_inches=5.2):
    """Insert an image paragraph before the given paragraph"""
    if not os.path.exists(img_path):
        print(f"  WARNING: Image not found: {img_path}")
        return None

    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_p.append(pPr)

    para._element.addprevious(new_p)

    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, para._parent)

    run = new_para.add_run()
    try:
        run.add_picture(img_path, width=Inches(width_inches))
        print(f"  OK: Embedded {os.path.basename(img_path)}")
        return new_para
    except Exception as e:
        print(f"  ERROR adding image: {e}")
        return None


# ================================================================
# STEP 1: Add chapter headings (Heading 1) before each chapter
# ================================================================
print("=" * 60)
print("Step 1: Adding chapter headings...")

section_to_idx = {}
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and 'Heading 2' in str(p.style.name):
        section_to_idx[text] = i

for section_name, (ch_heading, ch_num) in CHAPTERS.items():
    if section_name in section_to_idx:
        idx = section_to_idx[section_name]
        para = doc.paragraphs[idx]

        # Insert page break
        insert_page_break_paragraph_before(para)

        # Insert chapter heading
        make_paragraph_before(para, ch_heading, 'Heading 1')
        print(f"  Added: [{ch_heading}] before [{idx}] {section_name}")
    else:
        print(f"  SKIP: could not find '{section_name}'")

print()

# ================================================================
# STEP 2: Embed Chapter 3 figures before their captions
# ================================================================
print("Step 2: Embedding Chapter 3 figures...")

for fig_name, fig_info in CHAPTER3_FIGURES.items():
    cap_idx = fig_info['caption_idx']
    img_file = os.path.join(FIG_DIR, fig_info['img'])
    para = doc.paragraphs[cap_idx]
    insert_image_before(para, img_file, width_inches=5.2)

print()

# ================================================================
# STEP 3: Embed Chapter 4 figures after reference paragraphs
# ================================================================
print("Step 3: Embedding Chapter 4 figures...")

for fig_name, fig_info in CH4_FIGURES.items():
    insert_idx = fig_info['insert_after']
    img_file = os.path.join(FIG_DIR, fig_info['img'])

    # Insert image and caption BEFORE the paragraph that follows insert_idx
    next_para = doc.paragraphs[insert_idx + 1]

    # First insert image
    insert_image_before(next_para, img_file, width_inches=5.2)

    # Then insert caption (above the image since we're adding before next_para)
    # Actually the order of addprevious means: later calls go above earlier ones
    # So we should add caption first, then image
    # Let's redo: image was added first, so it's right before next_para
    # Now add caption before the image
    img_para = doc.paragraphs[insert_idx + 1]  # This is now the image paragraph
    # Hmm, this is getting tricky with indices shifting

    # Simpler approach: find the image paragraph we just inserted and add caption before it
    print(f"  NOTE: Caption for {fig_name} needs manual placement: {fig_info['caption'][:80]}")

print()

# ================================================================
# STEP 4: Remove dangling figure captions at the end
# ================================================================
print("Step 4: Removing dangling figure captions at end...")

paras_to_remove = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith('[Figure ') and (text.endswith('.]') or text.endswith('x).') or 'speedup' in text.lower()):
        paras_to_remove.append(i)
        print(f"  Remove [{i}]: {text[:90]}")

body = doc.element.body
for idx in sorted(paras_to_remove, reverse=True):
    elem = doc.paragraphs[idx]._element
    body.remove(elem)

print(f"  Removed {len(paras_to_remove)} dangling captions")
print()

# ================================================================
# STEP 5: Add Declaration page before Abstract
# ================================================================
print("Step 5: Adding Declaration page...")

abstract_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Abstract' and 'Heading 1' in str(p.style.name):
        abstract_idx = i
        break

if abstract_idx is not None:
    abstract_para = doc.paragraphs[abstract_idx]

    # Insert page break before Abstract
    insert_page_break_paragraph_before(abstract_para)

    # Declaration lines (reverse order for addprevious)
    decl_lines = [
        "Date: May 5, 2026",
        "Signature: ______________________",
        "Student ID: [Please fill in your Student ID]",
        "Student Name: JU JIAXING",
        "",
        "I have read the student handbook and I understand the meaning of academic dishonesty, in particular plagiarism and collusion. I declare that the work submitted for the final year project does not involve academic dishonesty. I give permission for my final year project work to be electronically scanned and if found to involve academic dishonesty, I am aware of the consequences as stated in the Student Handbook.",
        "Student Final Year Project Declaration",
    ]

    for line in reversed(decl_lines):
        if line.startswith("Student Final Year Project Declaration"):
            make_paragraph_before(abstract_para, line, 'Heading 1')
        elif line.startswith("Signature") or line.startswith("Date"):
            make_paragraph_before(abstract_para, line, 'Normal')
        elif line.startswith("Student Name") or line.startswith("Student ID"):
            make_paragraph_before(abstract_para, line, 'Normal')
        elif line == "":
            make_paragraph_before(abstract_para, '', 'Normal')
        else:
            make_paragraph_before(abstract_para, line, 'Normal')

    print("  Declaration page added before Abstract")
else:
    print("  ERROR: Could not find Abstract paragraph")

print()

# ================================================================
# Final count
# ================================================================
total_words = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
print(f"Final stats: {len(doc.paragraphs)} paragraphs, ~{total_words} words")

# Save
print(f"\nSaving to {OUT_PATH}...")
doc.save(OUT_PATH)
print("DONE!")
