import sys
sys.path.insert(0, r'C:\Users\16084\AppData\Roaming\Python\Python313\site-packages')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = r'C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing'
SRC_PATH = os.path.join(BASE_DIR, 'FYP_Thesis_Final.docx')
FIG_DIR = os.path.join(BASE_DIR, 'Figures')
OUT_PATH = os.path.join(BASE_DIR, 'FYP_Thesis_Final.docx')

doc = Document(SRC_PATH)


def make_paragraph_before(para, text='', style=None):
    new_p = OxmlElement('w:p')
    para._element.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, para._parent)
    if text:
        new_para.text = text
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def add_caption_after_image_before(next_para, caption_text):
    """Add caption between the image and next_para.
    Since image was added via addprevious, it sits before next_para.
    Adding another addprevious on next_para will put caption between image and next_para."""
    make_paragraph_before(next_para, caption_text, 'Normal')


# ================================================================
# Find where Ch4 images were inserted and add captions
# ================================================================
print("Step 3b: Adding captions for Chapter 4 figures...")

# We need to find the paragraphs that have images and add captions after them
# Strategy: scan paragraphs, find those near the target positions

# The original insert positions were: after paragraphs 341, 355, 362
# After inserting 6 chapter headings + page breaks = ~12 paragraphs added before Ch4
# So the indices shifted by roughly +12-15

# Let's find the figure paragraphs by scanning for paragraphs near the target areas
ch4_captions = [
    ('hello_e203', 'Figure 4.1: ILA capture showing PC progression during hello_e203 execution on FPGA.'),
    ('NICE', 'Figure 4.2: ILA capture showing NICE custom instruction activity on the FPGA.'),
    ('speedup', 'Figure 4.3: Performance comparison showing CNN accelerator speedup over software-only execution.'),
]

# Scan paragraphs for inline shapes (images)
img_paras = []
for i, p in enumerate(doc.paragraphs):
    # Check if paragraph has an image (drawing element)
    drawings = p._element.findall('.//' + qn('w:drawing'))
    if drawings:
        text_before = doc.paragraphs[i-1].text.strip()[:80] if i > 0 else ''
        text_after = doc.paragraphs[i+1].text.strip()[:80] if i+1 < len(doc.paragraphs) else ''
        img_paras.append((i, text_before, text_after))

print(f"Found {len(img_paras)} paragraphs with images:")
for idx, before, after in img_paras:
    print(f"  [{idx}] before=[{before}] after=[{after}]")

# Now determine which image is which based on surrounding text
# and add captions where needed
for idx, before, after in img_paras:
    # Check if there's already a caption after this image
    if after.startswith('Figure '):
        print(f"  [{idx}] Already has caption: {after[:60]}")
        continue

    # Determine which caption to use based on context
    if 'ILA capture' in before or 'PC progression' in before or 'hello_e203' in before:
        caption = 'Figure 4.1: ILA capture showing PC progression during hello_e203 execution on FPGA.'
    elif 'NICE' in before or 'nice' in before.lower():
        caption = 'Figure 4.2: ILA capture showing NICE custom instruction activity on the FPGA.'
    elif 'performance' in before.lower() or 'comparison' in before.lower():
        caption = 'Figure 4.3: Performance comparison showing CNN accelerator speedup over software-only execution.'
    else:
        # Fallback: look at the paragraph after the image
        caption = None
        print(f"  [{idx}] Unknown image context, before=[{before[:60]}], after=[{after[:60]}]")
        continue

    next_para = doc.paragraphs[idx + 1]
    add_caption_after_image_before(next_para, caption)
    print(f"  [{idx}] Added caption: {caption[:60]}")

print()

# ================================================================
# Also add the verification chain figure to Ch4 (Figure 4.4 or in discussion)
# ================================================================
print("Checking for fig_verification_chain.png...")
vchain_path = os.path.join(FIG_DIR, 'fig_verification_chain.png')
if os.path.exists(vchain_path):
    print(f"  Available: {vchain_path}")
else:
    print(f"  Not found")

print()

# ================================================================
# Print final structure summary
# ================================================================
print("Final heading structure:")
for i, p in enumerate(doc.paragraphs):
    style = str(p.style.name)
    if 'Heading' in style:
        indent = '  ' if 'Heading 2' in style else ''
        print(f"  [{i}] {indent}[{style}] {p.text[:90]}")

# Word count
total_words = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
print(f"\nTotal words: ~{total_words}")
print(f"Paragraphs: {len(doc.paragraphs)}")

# Save
print(f"\nSaving...")
doc.save(OUT_PATH)
print("DONE!")
