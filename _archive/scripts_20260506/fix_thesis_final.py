"""
One clean pass: start from original, remove dangling captions FIRST, then add everything.
"""
import sys
sys.path.insert(0, r'C:\Users\16084\AppData\Roaming\Python\Python313\site-packages')

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os

BASE_DIR = r'C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing'
SRC_PATH = os.path.join(BASE_DIR, 'FYP_Thesis_Formatted.docx')
FIG_DIR = os.path.join(BASE_DIR, 'Figures')
OUT_PATH = os.path.join(BASE_DIR, 'FYP_Thesis_Final.docx')

doc = Document(SRC_PATH)

print("=" * 60)
print("Starting from:", SRC_PATH)

# ================================================================
# STEP 1: Remove dangling figure captions at end FIRST
# ================================================================
print("\nStep 1: Removing dangling figure captions...")
paras_to_remove = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if (text.startswith('[Figure ') and
        ('System-level' in text or 'instruction format' in text or
         'microarchitecture' in text or 'PE array organization' in text or
         'Packed INT8' in text or 'build pipeline' in text or
         'ILA capture showing PC' in text or 'UART terminal output' in text)):
        paras_to_remove.append(i)
        print(f"  Remove [{i}]: {text[:90]}")

body = doc.element.body
for idx in sorted(paras_to_remove, reverse=True):
    body.remove(doc.paragraphs[idx]._element)
print(f"  Removed {len(paras_to_remove)} dangling captions")

# ================================================================
# STEP 2: Rebuild the index map on the CLEANED document
# ================================================================
print("\nStep 2: Building index map...")

h2_map = {}
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = str(p.style.name)
    if text and 'Heading 2' in style:
        h2_map[text] = i

# Find Abstract
abstract_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Abstract' and 'Heading 1' in str(p.style.name):
        abstract_idx = i
        break

print(f"  Abstract at [{abstract_idx}]")

# ================================================================
# STEP 3: Define all insertions (original indices from CLEANED doc)
# ================================================================

# Chapter headings
chapters = [
    ('1.1 Background', 'Chapter 1. Introduction'),
    ('2.1 RISC-V ISA and Custom Instruction Extensions', 'Chapter 2. Background'),
    ('3.1 System Architecture Overview', 'Chapter 3. Methodology'),
    ('4.1 RTL Simulation', 'Chapter 4. Results'),
    ('5.1 FPGA Bring-up Challenges and Resolution', 'Chapter 5. Discussion'),
    ('6.1 Summary of Contributions', 'Chapter 6. Conclusion'),
]

# Ch3 figures (insert image before caption paragraph)
fig3_list = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith('Figure 3.1:') and 'System-level' in text:
        fig3_list.append(('Figure 3.1', i, 'fig3_1_soc_architecture.png'))
    elif text.startswith('Figure 3.2:') and 'instruction format' in text:
        fig3_list.append(('Figure 3.2', i, 'fig3_2_instruction_format.png'))
    elif text.startswith('Figure 3.3:') and 'microarchitecture' in text:
        fig3_list.append(('Figure 3.3', i, 'fig3_3_pe_microarchitecture.png'))
    elif text.startswith('Figure 3.4:') and 'PE array' in text:
        fig3_list.append(('Figure 3.4', i, 'fig3_4_pe_array.png'))
    elif text.startswith('Figure 3.5:') and 'Packed' in text:
        fig3_list.append(('Figure 3.5', i, 'fig3_5_packed_format.png'))
    elif text.startswith('Figure 3.6:') and 'build pipeline' in text:
        fig3_list.append(('Figure 3.6', i, 'fig3_6_build_pipeline.png'))

# Ch4 figures (insert after specific paragraphs)
# Find insertion points by scanning for specific content
ch4_figs_to_insert = []

# Figure 4.1: near ILA PC trace discussion in 4.3
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = str(p.style.name)
    if '4.3' in text and 'hello_e203' in text and 'Heading 2' in style:
        for j in range(i+1, min(i+25, len(doc.paragraphs))):
            if 'ILA capture (1024' in doc.paragraphs[j].text:
                ch4_figs_to_insert.append(('fig_ila_pc_trace.png',
                    'Figure 4.1: ILA capture showing PC progression during hello_e203 execution on FPGA.',
                    j))
                print(f"  Figure 4.1: insert after [{j}]")
                break
        break

# Figure 4.2: near NICE ILA discussion in 4.5
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = str(p.style.name)
    if '4.5' in text and 'NICE' in text and 'Heading 2' in style:
        for j in range(i+1, min(i+30, len(doc.paragraphs))):
            if 'ILA capture' in doc.paragraphs[j].text and 'NICE' in doc.paragraphs[j].text:
                ch4_figs_to_insert.append(('fig_ila_nice_activity.png',
                    'Figure 4.2: ILA capture showing NICE custom instruction activity on the FPGA.',
                    j))
                print(f"  Figure 4.2: insert after [{j}]")
                break
        break

# Figure 4.3: near benchmark/speedup
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'speedup' in text and 'DEMO PASSED' in text:
        ch4_figs_to_insert.append(('fig_speedup_bar.png',
            'Figure 4.3: CNN accelerator benchmark results showing speedup over software-only execution.',
            i - 1))  # insert after the paragraph before this one
        print(f"  Figure 4.3: insert after [{i-1}]")
        break
else:
    # Fallback: find 4.5.2 ILA Evidence or similar
    for i, p in enumerate(doc.paragraphs):
        if '4.5 NICE' in p.text.strip() and 'Heading 2' in str(p.style.name):
            for j in range(i+1, min(i+30, len(doc.paragraphs))):
                if 'GPIOA LED' in doc.paragraphs[j].text:
                    ch4_figs_to_insert.append(('fig_speedup_bar.png',
                        'Figure 4.3: CNN accelerator benchmark results showing speedup over software-only execution.',
                        j))
                    print(f"  Figure 4.3 (fallback): insert after [{j}]")
                    break
            break

# Figure 3.7: verification chain
vchain_found = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if '3.5' in text and 'FPGA Bring-up' in text and 'Heading 2' in str(p.style.name):
        for j in range(i+1, min(i+20, len(doc.paragraphs))):
            if 'verification' in doc.paragraphs[j].text.lower():
                ch4_figs_to_insert.append(('fig_verification_chain.png',
                    'Figure 3.7: Multi-stage verification flow from RTL simulation through FPGA board validation.',
                    j))
                vchain_found = True
                print(f"  Figure 3.7: insert after [{j}]")
                break
        break

# ================================================================
# STEP 4: Combine all insertions and sort by index DESCENDING
# ================================================================
all_insertions = []

for section_name, ch_heading in chapters:
    if section_name in h2_map:
        all_insertions.append(('chapter', h2_map[section_name], (ch_heading,)))

for fig_name, idx, img_file in fig3_list:
    all_insertions.append(('fig3', idx, (img_file,)))

for img_file, caption, ref_idx in ch4_figs_to_insert:
    # Insert after ref_idx, so target is ref_idx+1
    all_insertions.append(('ch4', ref_idx + 1, (img_file, caption)))

# Declaration
if abstract_idx is not None:
    all_insertions.append(('decl', abstract_idx, None))

# Sort descending
all_insertions.sort(key=lambda x: x[1], reverse=True)

print(f"\nStep 4: Processing {len(all_insertions)} insertions (reverse order)...")

for item in all_insertions:
    action, target_idx, data = item

    if action == 'chapter':
        ch_heading = data[0]
        para = doc.paragraphs[target_idx]

        # Insert page break
        new_p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        new_p.append(r)
        para._element.addprevious(new_p)

        # Insert Heading 1
        new_p = OxmlElement('w:p')
        para._element.addprevious(new_p)
        np = Paragraph(new_p, para._parent)
        np.text = ch_heading
        try:
            np.style = 'Heading 1'
        except Exception:
            pass
        print(f"  Chapter: [{ch_heading}]")

    elif action == 'fig3':
        img_file = data[0]
        para = doc.paragraphs[target_idx]
        img_path = os.path.join(FIG_DIR, img_file)

        # Insert centered image
        new_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        new_p.append(pPr)
        para._element.addprevious(new_p)
        np = Paragraph(new_p, para._parent)
        run = np.add_run()
        if os.path.exists(img_path):
            run.add_picture(img_path, width=Inches(5.2))
            print(f"  Image: {img_file}")
        else:
            print(f"  SKIP (not found): {img_file}")

    elif action == 'ch4':
        img_file, caption = data
        if target_idx >= len(doc.paragraphs):
            print(f"  SKIP ch4: index {target_idx} out of range")
            continue
        para = doc.paragraphs[target_idx]
        img_path = os.path.join(FIG_DIR, img_file)

        # Insert caption first (will be below image since both use addprevious)
        new_p = OxmlElement('w:p')
        para._element.addprevious(new_p)
        np = Paragraph(new_p, para._parent)
        np.text = caption
        try:
            np.style = 'Normal'
        except Exception:
            pass

        # Insert image (will be above caption)
        new_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        new_p.append(pPr)
        para._element.addprevious(new_p)
        np = Paragraph(new_p, para._parent)
        run = np.add_run()
        if os.path.exists(img_path):
            run.add_picture(img_path, width=Inches(5.2))
            print(f"  Ch4: {img_file} + caption")
        else:
            print(f"  SKIP (not found): {img_file}")

    elif action == 'decl':
        para = doc.paragraphs[target_idx]

        # Page break
        new_p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        new_p.append(r)
        para._element.addprevious(new_p)

        # Declaration content (insert in reverse order since using addprevious)
        decl_texts = [
            ('Normal', 'Date: May 5, 2026'),
            ('Normal', 'Signature: ______________________'),
            ('Normal', 'Student ID: [Please fill in your Student ID]'),
            ('Normal', 'Student Name: JU JIAXING'),
            ('Normal', ''),
            ('Normal', 'I have read the student handbook and I understand the meaning of academic dishonesty, in particular plagiarism and collusion. I declare that the work submitted for the final year project does not involve academic dishonesty. I give permission for my final year project work to be electronically scanned and if found to involve academic dishonesty, I am aware of the consequences as stated in the Student Handbook.'),
            ('Heading 1', 'Student Final Year Project Declaration'),
        ]

        for style, text in decl_texts:
            new_p = OxmlElement('w:p')
            para._element.addprevious(new_p)
            np = Paragraph(new_p, para._parent)
            np.text = text
            try:
                np.style = style
            except Exception:
                pass

        print(f"  Declaration: added before Abstract")

# ================================================================
# STEP 5: Verify
# ================================================================
print("\n" + "=" * 60)
print("FINAL STRUCTURE:")
print("=" * 60)

img_count = 0
for i, p in enumerate(doc.paragraphs):
    style = str(p.style.name)
    if 'Heading' in style:
        indent = '  ' if 'Heading 2' in style else ''
        print(f"[{i:4d}] {indent}[{style}] {p.text[:90]}")
    drawings = p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
    if drawings:
        img_count += 1

total_words = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
ref_count = sum(1 for p in doc.paragraphs if p.text.strip() and p.text.strip()[0] == '[' and p.text.strip()[1].isdigit())

print(f"\nParagraphs: {len(doc.paragraphs)}")
print(f"Words: ~{total_words}")
print(f"Images: {img_count}")
print(f"References: {ref_count}")
print(f"Est. pages (1.5 spacing): ~{total_words / 250:.0f}")

doc.save(OUT_PATH)
print(f"\nSaved to: {OUT_PATH}")
print("DONE!")
