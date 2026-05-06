"""
Fix thesis document: add chapter headings, embed images, add declaration.
Processes insertions in reverse order to avoid index shifting issues.
"""
import sys
sys.path.insert(0, r'C:\Users\16084\AppData\Roaming\Python\Python313\site-packages')

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os

BASE_DIR = r'C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing'
SRC_PATH = os.path.join(BASE_DIR, 'FYP_Thesis_Formatted.docx')
FIG_DIR = os.path.join(BASE_DIR, 'Figures')
OUT_PATH = os.path.join(BASE_DIR, 'FYP_Thesis_Final.docx')

doc = Document(SRC_PATH)


def insert_page_break_before(para):
    """Insert a page break paragraph before the given paragraph."""
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    new_p.append(r)
    para._element.addprevious(new_p)


def insert_heading1_before(para, text):
    """Insert a Heading 1 paragraph before the given paragraph."""
    new_p = OxmlElement('w:p')
    para._element.addprevious(new_p)
    new_para = Paragraph(new_p, para._parent)
    new_para.text = text
    try:
        new_para.style = 'Heading 1'
    except Exception:
        pass
    return new_para


def insert_image_before(para, img_path, width_inches=5.2):
    """Insert a centered image paragraph before the given paragraph."""
    if not os.path.exists(img_path):
        print(f"  WARNING: Image not found: {os.path.basename(img_path)}")
        return None

    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_p.append(pPr)
    para._element.addprevious(new_p)

    new_para = Paragraph(new_p, para._parent)
    run = new_para.add_run()
    try:
        run.add_picture(img_path, width=Inches(width_inches))
        return new_para
    except Exception as e:
        print(f"  ERROR: {e}")
        return None


def insert_normal_before(para, text):
    """Insert a Normal paragraph before the given paragraph."""
    new_p = OxmlElement('w:p')
    para._element.addprevious(new_p)
    new_para = Paragraph(new_p, para._parent)
    new_para.text = text
    try:
        new_para.style = 'Normal'
    except Exception:
        pass
    return new_para


# ================================================================
# STEP 1: Build index map of ALL target paragraphs in ORIGINAL doc
# ================================================================
print("Building paragraph index map...")

# Find all Heading 2 paragraphs
h2_map = {}  # text -> index
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = str(p.style.name)
    if text and 'Heading 2' in style:
        h2_map[text] = i

# Find Abstract
abstract_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Abstract' and 'Heading 1' in str(p.style.name):
        abstract_idx = i
        break

print(f"Abstract at [{abstract_idx}]")
print(f"Found {len(h2_map)} Heading 2 sections")

# ================================================================
# STEP 2: Define all insertions with ORIGINAL indices
# All insertions are: (original_index, action_type, data)
# Process in REVERSE index order to avoid shifting issues
# ================================================================

Insertion = []

# Chapter headings + page breaks (before first section of each chapter)
chapter_first_sections = {
    'Chapter 1. Introduction': '1.1 Background',
    'Chapter 2. Background': '2.1 RISC-V ISA and Custom Instruction Extensions',
    'Chapter 3. Methodology': '3.1 System Architecture Overview',
    'Chapter 4. Results': '4.1 RTL Simulation',
    'Chapter 5. Discussion': '5.1 FPGA Bring-up Challenges and Resolution',
    'Chapter 6. Conclusion': '6.1 Summary of Contributions',
}

for ch_heading, section_name in chapter_first_sections.items():
    if section_name in h2_map:
        idx = h2_map[section_name]
        Insertion.append(('chapter', idx, (section_name, ch_heading)))

# Figure images (before caption paragraph)
# Map figure name -> (idx, img_file)
figure_captions = {}
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'Figure 3.1:' in text and text.startswith('Figure '):
        figure_captions['Figure 3.1'] = (i, 'fig3_1_soc_architecture.png')
    elif 'Figure 3.2:' in text and text.startswith('Figure '):
        figure_captions['Figure 3.2'] = (i, 'fig3_2_instruction_format.png')
    elif 'Figure 3.3:' in text and text.startswith('Figure '):
        figure_captions['Figure 3.3'] = (i, 'fig3_3_pe_microarchitecture.png')
    elif 'Figure 3.4:' in text and text.startswith('Figure '):
        figure_captions['Figure 3.4'] = (i, 'fig3_4_pe_array.png')
    elif 'Figure 3.5:' in text and text.startswith('Figure '):
        figure_captions['Figure 3.5'] = (i, 'fig3_5_packed_format.png')
    elif 'Figure 3.6:' in text and text.startswith('Figure '):
        figure_captions['Figure 3.6'] = (i, 'fig3_6_build_pipeline.png')

for fig_name, (idx, img_file) in figure_captions.items():
    Insertion.append(('figure', idx, (fig_name, img_file)))

# Ch4 figure references - find paragraphs that reference these figures
# and insert image+caption after them
ch4_fig_inserts = [
    ('hello_e203', 'ILA capture', 'fig_ila_pc_trace.png', 'Figure 4.1: ILA capture showing PC progression during hello_e203 execution on FPGA.'),
    ('NICE', 'custom instruction', 'fig_ila_nice_activity.png', 'Figure 4.2: ILA capture showing NICE custom instruction activity on the FPGA.'),
    ('cnn_accel', 'benchmark', 'fig_speedup_bar.png', 'Figure 4.3: CNN accelerator benchmark results showing speedup over software-only execution.'),
]

for keyword1, keyword2, img_file, caption in ch4_fig_inserts:
    # Find a paragraph in Ch4 that contains relevant keywords and is near a good insertion point
    found = False
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # Look for Heading 2 sections in Ch4
        if '4.3' in text and 'Heading 2' in str(p.style.name) and 'hello' in keyword1:
            # Insert after the section heading + a few paragraphs
            # Find the next ILA-related paragraph
            for j in range(i+1, min(i+20, len(doc.paragraphs))):
                if 'ILA' in doc.paragraphs[j].text and 'capture' in doc.paragraphs[j].text.lower():
                    Insertion.append(('ch4_figure', j, (img_file, caption)))
                    found = True
                    break
            break
        elif '4.5' in text and 'Heading 2' in str(p.style.name) and 'nice' in keyword1.lower():
            for j in range(i+1, min(i+20, len(doc.paragraphs))):
                if 'ILA' in doc.paragraphs[j].text:
                    Insertion.append(('ch4_figure', j, (img_file, caption)))
                    found = True
                    break
            break
    if not found:
        print(f"  WARNING: Could not find insertion point for {img_file} with '{keyword1}'")

# Declaration page (before Abstract)
if abstract_idx is not None:
    Insertion.append(('declaration', abstract_idx, None))

# Remove dangling figure captions at end
dangling_indices = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith('[') and 'Figure ' in text and (text.endswith('.]') or 'speedup' in text.lower()):
        dangling_indices.append(i)

if dangling_indices:
    print(f"Found {len(dangling_indices)} dangling figure captions to remove")

# ================================================================
# STEP 3: Sort all insertions by original index, DESCENDING
# ================================================================
# For remove operations, just note the indices
# For insert operations, sort descending to process end-to-start

insert_sorted = sorted(Insertion, key=lambda x: x[1], reverse=True)

print(f"\nProcessing {len(insert_sorted)} insertions (reverse order)...")

for item in insert_sorted:
    action = item[0]
    orig_idx = item[1]

    if action == 'chapter':
        section_name, ch_heading = item[2]
        para = doc.paragraphs[orig_idx]
        insert_page_break_before(para)
        insert_heading1_before(para, ch_heading)
        print(f"  Chapter: [{ch_heading}] before [{orig_idx}] {section_name}")

    elif action == 'figure':
        fig_name, img_file = item[2]
        para = doc.paragraphs[orig_idx]
        img_path = os.path.join(FIG_DIR, img_file)
        insert_image_before(para, img_path, width_inches=5.2)
        print(f"  Figure: {img_file} before [{orig_idx}] {fig_name}")

    elif action == 'ch4_figure':
        img_file, caption = item[2]
        # Insert after orig_idx (i.e., before orig_idx+1)
        next_idx = orig_idx + 1
        if next_idx < len(doc.paragraphs):
            next_para = doc.paragraphs[next_idx]
            img_path = os.path.join(FIG_DIR, img_file)
            insert_image_before(next_para, img_path, width_inches=5.2)
            insert_normal_before(next_para, caption)
            print(f"  Ch4 Fig: {img_file} + caption after [{orig_idx}]")
        else:
            print(f"  SKIP Ch4 Fig: {img_file} - index {orig_idx} is last paragraph")

    elif action == 'declaration':
        abstract_para = doc.paragraphs[orig_idx]
        insert_page_break_before(abstract_para)

        decl_lines = [
            ("Normal", "Date: May 5, 2026"),
            ("Normal", "Signature: ______________________"),
            ("Normal", "Student ID: [Please fill in your Student ID]"),
            ("Normal", "Student Name: JU JIAXING"),
            ("Normal", ""),
            ("Normal", "I have read the student handbook and I understand the meaning of academic dishonesty, in particular plagiarism and collusion. I declare that the work submitted for the final year project does not involve academic dishonesty. I give permission for my final year project work to be electronically scanned and if found to involve academic dishonesty, I am aware of the consequences as stated in the Student Handbook."),
            ("Heading 1", "Student Final Year Project Declaration"),
        ]

        for style, text in decl_lines:
            if style == 'Heading 1':
                insert_heading1_before(abstract_para, text)
            elif text == "":
                insert_normal_before(abstract_para, "")
            else:
                insert_normal_before(abstract_para, text)

        print(f"  Declaration: added before Abstract at [{orig_idx}]")

# ================================================================
# STEP 4: Remove dangling figure captions
# ================================================================
print(f"\nRemoving {len(dangling_indices)} dangling captions...")
body = doc.element.body
for idx in sorted(dangling_indices, reverse=True):
    elem = doc.paragraphs[idx]._element
    body.remove(elem)

# ================================================================
# STEP 5: Verify final structure
# ================================================================
print("\n" + "=" * 60)
print("FINAL DOCUMENT STRUCTURE:")
print("=" * 60)

for i, p in enumerate(doc.paragraphs):
    style = str(p.style.name)
    if 'Heading' in style:
        indent = '  ' if 'Heading 2' in style else ''
        print(f"[{i:4d}] {indent}[{style}] {p.text[:90]}")

total_words = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
print(f"\nParagraphs: {len(doc.paragraphs)}")
print(f"Words: ~{total_words}")

# Save
print(f"\nSaving to {OUT_PATH}...")
doc.save(OUT_PATH)
print("DONE! Open in Word to verify formatting.")
