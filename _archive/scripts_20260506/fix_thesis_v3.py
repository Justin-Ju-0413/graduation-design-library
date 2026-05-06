"""
Final cleanup: remove dangling captions, add references heading, add missing figure
"""
import sys
sys.path.insert(0, r'C:\Users\16084\AppData\Roaming\Python\Python313\site-packages')

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os

BASE_DIR = r'C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing'
SRC_PATH = os.path.join(BASE_DIR, 'FYP_Thesis_Final.docx')
FIG_DIR = os.path.join(BASE_DIR, 'Figures')
OUT_PATH = SRC_PATH

doc = Document(SRC_PATH)


def insert_normal_before(para, text):
    new_p = OxmlElement('w:p')
    para._element.addprevious(new_p)
    new_para = Paragraph(new_p, para._parent)
    new_para.text = text
    try:
        new_para.style = 'Normal'
    except Exception:
        pass
    return new_para


def insert_image_before(para, img_path, width_inches=5.2):
    if not os.path.exists(img_path):
        print(f"  WARNING: Image not found: {img_path}")
        return None
    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_p.append(pPr)
    para._element.addprevious(new_p)
    new_para = Paragraph(new_p, para._parent)
    run = new_para.add_run()
    try:
        run.add_picture(img_path, width=Inches(width_inches))
        return new_para
    except Exception as e:
        print(f"  ERROR: {e}")
        return None


# ================================================================
# STEP 1: Remove dangling figure captions at end
# ================================================================
print("Step 1: Removing dangling figure captions...")

paras_to_remove = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    # Match: [Figure X.Y: ...]  (these are the dangling ones after references)
    if (text.startswith('[Figure ') and
        ('System-level' in text or 'instruction format' in text or
         'microarchitecture' in text or 'PE array organization' in text or
         'Packed INT8' in text or 'build pipeline' in text or
         'ILA capture showing' in text or 'UART terminal' in text)):
        paras_to_remove.append(i)
        print(f"  Remove [{i}]: {text[:90]}")

body = doc.element.body
for idx in sorted(paras_to_remove, reverse=True):
    elem = doc.paragraphs[idx]._element
    body.remove(elem)

print(f"  Removed {len(paras_to_remove)} dangling captions\n")

# ================================================================
# STEP 2: Add References heading
# ================================================================
print("Step 2: Adding References heading...")

# Find the first reference paragraph [1]
ref1_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('[1] '):
        ref1_idx = i
        break

if ref1_idx:
    para = doc.paragraphs[ref1_idx]
    insert_normal_before(para, '')  # blank line
    # Insert as Heading 1
    new_p = OxmlElement('w:p')
    para._element.addprevious(new_p)
    new_para = Paragraph(new_p, para._parent)
    new_para.text = 'References'
    try:
        new_para.style = 'Heading 1'
    except Exception:
        pass
    print(f"  Added 'References' heading before [{ref1_idx}]\n")
else:
    print("  Could not find [1] reference\n")

# ================================================================
# STEP 3: Add fig_speedup_bar.png in Ch4.5 area
# ================================================================
print("Step 3: Adding fig_speedup_bar.png...")

# Find a good spot - near discussion of benchmark results in 4.5 or Discussion
speedup_img = os.path.join(FIG_DIR, 'fig_speedup_bar.png')

# Look for paragraph about benchmark/speedup near 4.5 NICE section
target_idx = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = str(p.style.name)
    if '4.5' in text and 'Heading 2' in style:
        # Found 4.5 section, look for LED/GPIO or PASS/FAIL discussion
        for j in range(i+1, min(i+30, len(doc.paragraphs))):
            t = doc.paragraphs[j].text.strip()
            if 'PASS' in t or 'GPIOA LED' in t:
                target_idx = j
                break
        break

if target_idx:
    next_para = doc.paragraphs[target_idx + 1]
    insert_image_before(next_para, speedup_img, width_inches=4.8)
    insert_normal_before(next_para,
        'Figure 4.3: CNN accelerator on-board test results showing DEMO PASSED status with GPIO LED indication.')
    print(f"  Embedded fig_speedup_bar.png after [{target_idx}]\n")
else:
    print("  Could not find insertion point for speedup bar\n")

# ================================================================
# STEP 4: Add fig_verification_chain.png (in Ch3 or Ch4)
# ================================================================
print("Step 4: Adding fig_verification_chain.png...")

vchain_img = os.path.join(FIG_DIR, 'fig_verification_chain.png')

# Find multi-stage validation discussion in Ch3.5 or early Ch4
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'multi-stage validation' in text.lower() or 'three-stage verification' in text.lower():
        next_para = doc.paragraphs[i + 1]
        insert_image_before(next_para, vchain_img, width_inches=5.0)
        insert_normal_before(next_para,
            'Figure 3.7: Multi-stage verification flow from RTL simulation through FPGA board validation.')
        print(f"  Embedded fig_verification_chain.png after [{i}]\n")
        break
else:
    # Fallback: add in Ch3.5 FPGA Bring-up Methodology area
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = str(p.style.name)
        if '3.5' in text and 'Heading 2' in style:
            for j in range(i+1, min(i+20, len(doc.paragraphs))):
                t = doc.paragraphs[j].text.strip()
                if 'verification' in t.lower() and 'methodology' in t.lower():
                    next_para = doc.paragraphs[j + 1]
                    insert_image_before(next_para, vchain_img, width_inches=5.0)
                    insert_normal_before(next_para,
                        'Figure 3.7: Multi-stage verification flow from RTL simulation through FPGA board validation.')
                    print(f"  Embedded fig_verification_chain.png after [{j}] (fallback)\n")
                    break
            break

# ================================================================
# STEP 5: Final summary
# ================================================================
print("=" * 60)
print("FINAL HEADING STRUCTURE:")
print("=" * 60)

img_count = 0
for i, p in enumerate(doc.paragraphs):
    style = str(p.style.name)
    if 'Heading' in style:
        indent = '  ' if 'Heading 2' in style else ''
        print(f"[{i:4d}] {indent}[{style}] {p.text[:90]}")
    drawings = p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
    if drawings:
        img_count += 1

total_words = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print(f"Total words: ~{total_words}")
print(f"Embedded images: {img_count}")
print(f"Estimated pages (250w/p, 1.5 spacing): ~{total_words / 250:.0f}")

doc.save(OUT_PATH)
print(f"\nSaved to {OUT_PATH}")
print("DONE! Ready for Word manual review.")
