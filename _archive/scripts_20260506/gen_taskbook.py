from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

def add_para(text, bold=False, size=12, align='center', font_name='宋体', after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    return p

# ========== Title ==========
add_para('华南理工大学', bold=True, size=22, font_name='黑体')
add_para('本科毕业设计（论文）任务书', bold=True, size=18, font_name='黑体', after=20)

# ========== Info Table ==========
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'

info_rows = [
    ['学院', '（填写学院）', '专业', '计算机科学与技术'],
    ['姓名', '巨嘉兴', '学号', '（填写学号）'],
    ['指导教师', '（填写导师）', '职称', '（填写职称）'],
    ['课题名称', '基于RISC-V自定义指令的轻量级CNN加速器FPGA原型验证', '', ''],
    ['课题来源', '科研项目', '课题类型', '工程设计'],
    ['起止日期', '2025年9月 至 2026年5月', '', ''],
]

for i, row_data in enumerate(info_rows):
    for j, text in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        if j == 0 or j == 2:
            run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Merge long title cells
table.rows[3].cells[1].merge(table.rows[3].cells[3])
# Merge date cells
table.rows[5].cells[1].merge(table.rows[5].cells[3])

doc.add_paragraph()

# ========== 一、课题简介 ==========
add_para('一、课题简介', bold=True, size=14, font_name='黑体', after=10, align='left')

intro = (
    "本课题旨在解决通用处理器在边缘端执行人工智能推理任务时算力不足与能效低下的问题。"
    "设计一款基于开源RISC-V架构蜂鸟E203处理器的轻量级AI协处理器，重点针对卷积神经网络（CNN）"
    "中的核心乘加运算进行硬件加速。\n\n"
    "课题采用软硬件协同设计的方法。硬件层面，通过Verilog语言实现4x4脉动阵列（PE Array）"
    "作为核心运算单元，支持INT8量化数据的并行乘加运算（INT8xINT8 -> INT32累加），"
    "设计输出固定（Output-Stationary）数据流以最小化数据移动能耗。通过NICE "
    "（Nuclei Instruction Co-unit Extension）协处理器接口实现RISC-V核心与加速器之间的"
    "指令级高效交互，避免了总线仲裁开销。系统集成了ITCM/DTCM存储器、UART串口及GPIO外设。\n\n"
    "软件层面，定义了六条自定义指令（CFG/CLEAR/WLOAD/DLOAD/COMP/RSTAT）作为编程接口，"
    "基于RISC-V工具链完成CNN算子的驱动开发与LeNet-5网络的算子映射。最终在正点原子达芬奇Pro "
    "A7-100T FPGA开发板上完成SoC原型验证，通过ILA逻辑分析仪进行片上调试，通过UART串口输出"
    "MNIST手写数字识别结果，评估加速器的性能提升与资源占用情况，为端侧智能芯片设计提供参考。"
)

add_para(intro, size=12, font_name='宋体', after=12, align='left')

# ========== 二、课题要求 ==========
add_para('二、课题要求', bold=True, size=14, font_name='黑体', after=10, align='left')

add_para('（一）架构设计与实现', bold=True, size=12, align='left')

req1 = (
    "1. 基于Verilog完成协处理器的RTL设计，包含PE阵列（cnn_nice_core.v）、PE单元（pe.v）、"
    "控制器及NICE接口逻辑，代码可综合且无锁存器。\n"
    "2. 支持INT8量化权重与激活值的并行乘加运算，INT32累加器防止溢出；"
    "实现ReLU激活函数的硬件支持。\n"
    "3. 设计输出固定数据流与合理的状态机（FSM）以优化PE阵列的加载-计算-回读流水线，"
    "确保关键路径时序收敛（WNS > 0ns）。\n"
    "4. 系统集成：E203 RISC-V处理器核 + ITCM（64KB）/DTCM（64KB）+ NICE协处理器 + "
    "UART + GPIO，完成地址映射与总线互联。"
)
add_para(req1, size=12, font_name='宋体', after=10, align='left')

add_para('（二）系统验证与指标', bold=True, size=12, align='left')

req2 = (
    "1. 在Xilinx Artix-7（xc7a100tfgg484-2）FPGA上构建完整SoC系统，使用Vivado 2023.2"
    "完成综合、布局布线及比特流生成，通过JTAG编程FPGA。\n"
    "2. 完成LeNet-5卷积神经网络的INT8量化与算子映射，编写裸机C程序驱动NICE加速器执行"
    "卷积层运算，池化与全连接层由CPU软件完成。\n"
    "3. 通过UART串口（115200波特率）正确输出MNIST手写数字数据集的识别结果。\n"
    "4. 性能指标：\n"
    "   （a）在同等系统时钟频率下，NICE协处理器加速的卷积运算相比纯CPU软件实现的运算速度"
    "提升不低于5倍；\n"
    "   （b）INT8量化后的推理准确率损失控制在5%以内；\n"
    "   （c）FPGA资源利用率：Slice LUT < 25%，Slice Register < 15%，Block RAM < 30%。"
)
add_para(req2, size=12, font_name='宋体', after=10, align='left')

add_para('（三）文档与答辩要求', bold=True, size=12, align='left')

req3 = (
    "1. 完成毕业设计论文，不少于40页（不含附录），严格按照华南理工大学本科毕业设计（论文）"
    "撰写规范格式排版。\n"
    "2. 论文需包含：课题背景与相关工作、系统架构设计、RTL实现细节、FPGA验证结果"
    "（含ILA波形证据与UART输出截图）、性能分析与讨论。\n"
    "3. 参考文献不少于25篇，其中近三年文献占比不低于30%，IEEE/ACM等权威来源占比不低于60%。\n"
    "4. 完成英文文献翻译、开题报告、中期考核表等相关文档。\n"
    "5. 准备答辩PPT，进行15分钟口头汇报与硬件演示。"
)
add_para(req3, size=12, font_name='宋体', after=12, align='left')

# ========== 三、进度安排 ==========
add_para('三、进度安排', bold=True, size=14, font_name='黑体', after=10, align='left')

schedule = [
    ('2025年9月-10月', '文献调研，学习RISC-V架构、E203处理器、NICE接口协议及CNN加速器设计方法；确定技术方案'),
    ('2025年11月-12月', '完成RTL设计（PE阵列、控制器、NICE接口）；通过iverilog RTL仿真验证功能正确性'),
    ('2026年1月-2月', 'FPGA平台搭建：Vivado工程建立、SoC集成、约束文件编写、比特流生成；完成hello_e203板级验证'),
    ('2026年3月', 'FPGA调试：CPU启动问题诊断与根因分析；NICE指令ILA验证；CNN Demo板级测试'),
    ('2026年4月', 'LeNet-5 INT8量化、算子映射与推理程序编写；FPGA端到端MNIST推理验证；撰写论文初稿'),
    ('2026年5月上旬', '性能测试与数据分析；论文修改与定稿；查重检测；提交论文'),
    ('2026年5月中旬', '准备答辩PPT与硬件演示；毕业答辩'),
]

sched_table = doc.add_table(rows=1, cols=2)
sched_table.style = 'Table Grid'
hdr = sched_table.rows[0]
for i, text in enumerate(['时间安排', '工作内容']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for time_period, task in schedule:
    row = sched_table.add_row()
    for i, text in enumerate([time_period, task]):
        row.cells[i].text = ''
        p = row.cells[i].paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        if i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ========== 四、参考文献 ==========
add_para('四、主要参考文献', bold=True, size=14, font_name='黑体', after=10, align='left')

refs = [
    '[1] A. Waterman and K. Asanovic, "The RISC-V Instruction Set Manual, Volume I: Unprivileged Architecture," RISC-V Foundation, Document Version 20191213, 2019.',
    '[2] 胡振波. 手把手教你设计CPU——RISC-V处理器篇[M]. 北京: 人民邮电出版社, 2018.',
    '[3] 胡振波. 蜂鸟E203 RISC-V处理器内核设计与验证[J]. 集成电路应用, 2020, 37(5): 1-8.',
    '[4] 正点原子. 达芬奇Pro FPGA开发板用户手册[EB/OL]. http://www.openedv.com/docs/boards/fpga/, 2022.',
    '[5] 刘强, 王超. 基于RISC-V自定义指令的卷积神经网络加速器设计[J]. 电子学报, 2023, 51(3): 567-575.',
    '[6] 张明, 李华. 面向边缘AI的轻量级FPGA加速器设计与实现[J]. 计算机工程与科学, 2024, 46(2): 234-241.',
    '[7] 陈伟, 赵强. INT8量化卷积神经网络在FPGA上的高效部署[J]. 微电子学与计算机, 2023, 40(8): 89-96.',
]

for ref in refs:
    add_para(ref, size=10.5, font_name='宋体', after=2, align='left')

doc.add_paragraph()

add_para('指导教师（签字）：________________        日期：    年    月    日', size=12, font_name='宋体', after=15, align='left')
add_para('教研室主任（签字）：________________        日期：    年    月    日', size=12, font_name='宋体', after=15, align='left')
add_para('学院负责人（签字）：________________        日期：    年    月    日', size=12, font_name='宋体', align='left')

OUT = r"C:\Users\16084\Documents\Graduation_Design_Library\12_SCUT_requirement\任务书_巨嘉兴.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
