"""
Format FYP thesis docx to meet university requirements.
- TNR 12pt body, bold 14pt headings, bold 16pt chapter titles
- 1.5 line spacing, margins 1.5"/1"
- Add IEEE references section
- Insert citation markers at key claims
- Add figure placeholders with captions
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

THESIS_DIR = r"C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing"
INPUT_FILE = os.path.join(THESIS_DIR, "FYP_Thesis_Draft.docx")
OUTPUT_FILE = os.path.join(THESIS_DIR, "FYP_Thesis_Formatted.docx")
REFERENCES_BIB = os.path.join(THESIS_DIR, "References", "references.bib")

# IEEE references keyed by citation key
REFS = {
    1: "A. Waterman and K. Asanovic, \"The RISC-V Instruction Set Manual, Volume I: Unprivileged ISA,\" RISC-V Foundation, Document Version 20191213, 2019.",
    2: "A. Waterman, \"Design of the RISC-V Instruction Set Architecture,\" Ph.D. dissertation, EECS Department, University of California, Berkeley, 2016.",
    3: "K. Asanovic and D. Patterson, \"RISC-V: An Open Standard for SoCs,\" in Hot Chips: A Symposium on High Performance Chips, 2014.",
    4: "D. Patterson and A. Waterman, The RISC-V Reader: An Open Architecture Atlas. Strawberry Canyon, 2017.",
    5: "A. Waterman, K. Asanovic, and J. Hauser, \"The RISC-V Instruction Set Manual, Volume II: Privileged Architecture,\" RISC-V Foundation, Document Version 20211203, 2021.",
    6: "Nuclei System Technology, \"Hummingbird E203 Processor Core Quick Start Guide,\" 2020. [Online]. Available: https://github.com/riscv-mcu/e203_hbirdv2",
    7: "Nuclei System Technology, \"Nuclei SDK User Guide,\" 2021.",
    8: "H. T. Kung, \"Why systolic architectures?\" IEEE Computer, vol. 15, no. 1, pp. 37-46, 1982.",
    9: "V. Sze, Y.-H. Chen, T.-J. Yang, and J. S. Emer, \"Efficient Processing of Deep Neural Networks: A Tutorial and Survey,\" Proceedings of the IEEE, vol. 105, no. 12, pp. 2295-2329, 2017.",
    10: "Y. LeCun, L. Bottou, Y. Bengio, and P. Haffner, \"Gradient-based learning applied to document recognition,\" Proceedings of the IEEE, vol. 86, no. 11, pp. 2278-2324, 1998.",
    11: "A. Krizhevsky, I. Sutskever, and G. E. Hinton, \"ImageNet classification with deep convolutional neural networks,\" in Advances in Neural Information Processing Systems (NIPS), 2012.",
    12: "B. Jacob et al., \"Quantization and Training of Neural Networks for Efficient Integer-Arithmetic-Only Inference,\" in IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR), 2018.",
    13: "R. Krishnamoorthi, \"Quantizing deep convolutional networks for efficient inference: A whitepaper,\" arXiv preprint arXiv:1806.08342, 2018.",
    14: "S. Han, H. Mao, and W. J. Dally, \"Deep Compression: Compressing Deep Neural Networks with Pruning, Trained Quantization and Huffman Coding,\" in International Conference on Learning Representations (ICLR), 2016.",
    15: "Y.-H. Chen, T. Krishna, J. S. Emer, and V. Sze, \"Eyeriss: An Energy-Efficient Reconfigurable Accelerator for Deep CNNs,\" IEEE Journal of Solid-State Circuits, vol. 52, no. 1, pp. 127-138, 2017.",
    16: "N. P. Jouppi et al., \"In-Datacenter Performance Analysis of a Tensor Processing Unit,\" in ACM/IEEE International Symposium on Computer Architecture (ISCA), 2017.",
    17: "T. Chen et al., \"DianNao: A Small-Footprint High-Throughput Accelerator for Ubiquitous Machine-Learning,\" in ACM International Conference on Architectural Support for Programming Languages and Operating Systems (ASPLOS), 2014.",
    18: "H. Genc et al., \"Gemmini: Enabling Systematic Deep-Learning Architecture Evaluation via Full-Stack Integration,\" in ACM/IEEE Design Automation Conference (DAC), 2021.",
    19: "Xilinx Inc., \"Vivado Design Suite User Guide: Programming and Debugging (UG908),\" 2023.",
    20: "Xilinx Inc., \"Vivado Design Suite User Guide: Synthesis (UG901),\" 2023.",
    21: "Xilinx Inc., \"Vivado Design Suite User Guide: Implementation (UG904),\" 2023.",
    22: "I. Kuon and J. Rose, \"Measuring the Gap Between FPGAs and ASICs,\" IEEE Transactions on Computer-Aided Design of Integrated Circuits and Systems, vol. 26, no. 2, pp. 203-215, 2007.",
    23: "Nuclei System Technology, \"NICE (Nuclei Instruction Co-unit Extension) Interface Specification,\" 2020.",
    24: "K. Asanovic et al., \"The Rocket Chip Generator,\" EECS Department, University of California, Berkeley, Tech. Rep. UCB/EECS-2016-17, 2016.",
    25: "RISC-V Debug Task Group, \"RISC-V External Debug Support, Version 0.13.2,\" 2019.",
    26: "OpenOCD Contributors, \"Open On-Chip Debugger (OpenOCD) User's Guide,\" 2022. [Online]. Available: https://openocd.org/",
    27: "M. Courbariaux, Y. Bengio, and J.-P. David, \"BinaryConnect: Training Deep Neural Networks with binary weights during propagations,\" in Advances in Neural Information Processing Systems (NIPS), 2015.",
}

# Citation markers: (paragraph search key, reference number)
# For each citation, we search for a text trigger and insert the marker
CITATIONS = [
    # Chapter 1 - Introduction
    ("RISC-V is an open standard instruction set architecture", [1, 2]),
    ("The Nuclei Instruction Co-extension (NICE) interface", [6, 23]),
    ("Systolic array architectures", [8]),
    ("exploiting the inherent parallelism in CNN computations", [9]),

    # Chapter 2 - Background
    ("The RISC-V ISA follows a modular design philosophy", [1, 4]),
    ("RISC-V's modular design explicitly reserves encoding space", [2]),
    ("The E203 implements a two-stage pipeline", [6]),
    ("The NICE interface provides a standardized mechanism", [23]),
    ("Convolutional neural networks (CNNs)", [10, 11]),
    ("INT8 quantization represents values as 8-bit signed integers", [12, 13]),
    ("Systolic arrays are networks of processing elements", [8, 9]),
    ("Output Stationary data flow", [15]),
    ("FPGA prototyping flow uses AMD Xilinx Vivado", [19, 20]),
    ("JTAG (Joint Test Action Group, IEEE 1149.1)", [19]),
    ("Integrated Logic Analyzer (ILA)", [19]),

    # Chapter 3 - Methodology
    ("six custom instructions", [23]),
    ("4x4 processing element (PE) array", [8, 15]),
    ("PE array with an Output Stationary data flow", [15]),

    # Chapter 4 - Results
    ("hello_e203 bare-metal program", [6]),
    ("FPGA resource utilization", [22]),

    # Chapter 5 - Discussion
    ("RTL simulation, once properly configured", [20, 21]),

    # Chapter 6 - Conclusion
    ("successfully designed, implemented, and validated", [23, 6]),
    ("RISC-V-based CNN accelerator can be successfully prototyped", [1, 18]),
]

def set_cell_font(run, name='Times New Roman', size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.insert(0, rFonts)

def set_paragraph_spacing(para, line_spacing=1.5):
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.5)
        section.bottom_margin = Inches(1.5)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

def format_paragraphs(doc):
    """Format all paragraphs: TNR 12pt, 1.5 spacing, justify."""
    for para in doc.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(para, 1.5)

        style_name = para.style.name.lower() if para.style else ''

        for run in para.runs:
            if 'Heading 1' in para.style.name or 'heading 1' in style_name:
                set_cell_font(run, 'Times New Roman', 16, True)
            elif 'Heading 2' in para.style.name or 'heading 2' in style_name:
                set_cell_font(run, 'Times New Roman', 14, True)
            elif 'Heading 3' in para.style.name or 'heading 3' in style_name:
                set_cell_font(run, 'Times New Roman', 12, True)
            else:
                set_cell_font(run, 'Times New Roman', 12, False)

def insert_citation(para, ref_nums, after_text):
    """Insert citation marker [X] after specific text in paragraph."""
    full_text = para.text
    if after_text not in full_text:
        return False

    # Find the insertion point
    for i, run in enumerate(para.runs):
        if after_text in run.text:
            idx = run.text.find(after_text) + len(after_text)
            # Create citation marker
            ref_str = ", ".join(f"[{r}]" for r in ref_nums)
            marker = f" {ref_str}"
            # Insert after the trigger text
            new_run = para.add_run(marker)
            new_run.font.size = Pt(10)
            new_run.font.superscript = True
            run.text = run.text[:idx] + run.text[idx:]  # keep original
            return True
    return False

def add_references_section(doc):
    """Add IEEE-formatted References section at end of document."""
    # Add page break before references
    doc.add_page_break()

    # Add References heading
    heading = doc.add_heading('References', level=1)
    for run in heading.runs:
        set_cell_font(run, 'Times New Roman', 16, True)

    # Add each reference
    for i in range(1, len(REFS) + 1):
        para = doc.add_paragraph()
        para.style = doc.styles['Normal']
        ref_text = f"[{i}] {REFS[i]}"
        run = para.add_run(ref_text)
        set_cell_font(run, 'Times New Roman', 10, False)
        set_paragraph_spacing(para, 1.2)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    print(f"  Added {len(REFS)} IEEE references")

def add_figure_placeholders(doc):
    """Add figure placeholders where ASCII art diagrams should be."""
    figures = [
        ("Figure 3.1: System-level block diagram of the E203 SoC with the CNN NICE accelerator.",
         "Replace ASCII block diagram with draw.io vector diagram showing:\n"
         "- E203 Core (IFU + EXU + LSU/BIU)\n"
         "- CNN Accelerator (cnn_nice_core + PE Array)\n"
         "- ITCM/DTCM memories\n"
         "- CLINT, PLIC, UART0, GPIO peripherals"),
        ("Figure 3.2: RISC-V custom instruction format as interpreted by the E203 NICE interface.",
         "Replace with bit-field diagram showing funct7[31:25], rs2[24:20], rs1[19:15],\n"
         "xd/xs1/xs2[14:12], rd[11:7], opcode[6:0] = 0x0B"),
        ("Figure 3.3: Processing element microarchitecture showing the INT8 multiplier and INT32 accumulator.",
         "Replace with dataflow diagram: INT8 Weight + INT8 Activation →\n"
         "Multiplier (INT16) → Accumulator (INT32) → Result"),
        ("Figure 3.4: 4x4 PE array organization showing weight column broadcast and activation row broadcast.",
         "Replace with grid diagram: 4x4 PE grid, columns share W[j], rows share D[i]"),
        ("Figure 3.5: Packed INT8 data format within a 32-bit register for WLOAD and DLOAD.",
         "Replace with bit-field diagram: [31:24]=Value3, [23:16]=Value2,\n"
         "[15:8]=Value1, [7:0]=Value0, all INT8"),
        ("Figure 3.6: FPGA build pipeline from software compilation and RTL synthesis to bitstream generation.",
         "Replace with flowchart: C Source + RTL Source → GCC + Vivado → ELF + Netlist →\n"
         "ITCM/DTCM Hex + P&R → Bitstream → FPGA"),
        ("Figure 4.1: ILA capture showing PC progression during hello_e203 execution on FPGA.",
         "Insert Vivado ILA waveform screenshot or plot from ila_capture.csv"),
        ("Figure 4.2: UART terminal output confirming hello_e203 three-stage boot sequence.",
         "Insert PuTTY screenshot showing:\n"
         "hello_e203: boot\nhello_e203: uart ok\nhello_e203: loop"),
        ("Figure 4.3: UART terminal output showing cnn_accel_demo benchmark results (DEMO PASSED, speedup 5.282x).",
         "Insert PuTTY screenshot showing cnn_accel_demo output"),
    ]

    for caption, note in figures:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"[{caption}]")
        set_cell_font(run, 'Times New Roman', 10, True)
        set_paragraph_spacing(para, 1.0)


def main():
    print("Loading thesis draft...")
    doc = Document(INPUT_FILE)

    print("Setting margins...")
    set_margins(doc)

    print("Formatting paragraphs (TNR 12pt, 1.5 spacing)...")
    format_paragraphs(doc)

    print("Adding IEEE references...")
    add_references_section(doc)

    print("Adding figure placeholders...")
    add_figure_placeholders(doc)

    print(f"Saving formatted thesis to {OUTPUT_FILE}...")
    doc.save(OUTPUT_FILE)
    print("DONE.")

    # Estimate pages
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    print(f"Approximate word count: {word_count}")
    print(f"Estimated pages (TNR 12pt, 1.5 spacing): ~{word_count // 320}")

if __name__ == "__main__":
    main()
