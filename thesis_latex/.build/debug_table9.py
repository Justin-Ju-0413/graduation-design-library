import json, os, re
from docx import Document

DOCX_PATH = r'C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing\FYP_Thesis_Final_v2.docx'
doc = Document(DOCX_PATH)

LATEX_SPECIAL = {'\\': '\\textbackslash{}', '{': '\\{', '}': '\\}',
                  '$': '\\$', '&': '\\&', '#': '\\#', '_': '\\_',
                  '%': '\\%', '~': '\\textasciitilde{}', '^': '\\textasciicircum{}'}

def escape_latex(text):
    parts = re.split(r'(\$[^$]*\$)', text)
    result = []
    for part in parts:
        if part.startswith('$') and part.endswith('$'):
            result.append(part)
        else:
            for old, new in LATEX_SPECIAL.items():
                part = part.replace(old, new)
            result.append(part)
    return ''.join(result)

def extract_table_latex(table_idx):
    if table_idx >= len(doc.tables):
        return None, 0
    table = doc.tables[table_idx]
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    if not rows:
        return None, 0
    ncols = max(len(r) for r in rows)
    col_fmt = 'l' + 'c' * (ncols - 1) if ncols > 1 else 'l'
    latex_rows = []
    for idx, row in enumerate(rows):
        while len(row) < ncols:
            row.append('')
        escaped = [escape_latex(c) for c in row]
        line = ' & '.join(escaped) + ' \\\\'
        if idx == 0:
            line += ' \\hline\\hline'
        else:
            line += ' \\hline'
        latex_rows.append(line)
    tabular = '\\begin{tabular}{' + col_fmt + '}\n'
    tabular += '\\hline\n'
    tabular += '\n'.join(latex_rows) + '\n'
    tabular += '\\end{tabular}'
    return tabular, ncols

tabular, ncols = extract_table_latex(9)
print(f'tabular is None: {tabular is None}')
print(f'ncols: {ncols}')
if tabular:
    print(tabular[:300])
