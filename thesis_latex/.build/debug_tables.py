"""Debug script: table 9 extraction and chapter 4 table placement."""
import json, os, re
from docx import Document

DOCX_PATH = r"C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing\FYP_Thesis_Final_v2.docx"
doc = Document(DOCX_PATH)

with open(r'C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex\docx_structure.json', 'r', encoding='utf-8') as f:
    data = json.load(f)
paragraphs = data['paragraphs']

# Check table 9
print(f"Total tables: {len(doc.tables)}")
print(f"Table 9 rows: {len(doc.tables[9].rows)}")
for row in doc.tables[9].rows:
    print(f"  Row: {[c.text.strip()[:30] for c in row.cells]}")

# Check what paragraph 351 is in paragraphs
p351 = paragraphs[351]
print(f"\nPara 351: style={p351['style']} text='{p351['text'][:80]}'")

# Check chapter 4 range
# Chapter 4 starts at Heading 1 index 332
ch4_start = 333  # first content para
ch4_end = 404    # Chapter 5 starts at 404
print(f"\nChapter 4 range: {ch4_start} to {ch4_end}")
print(f"Para 351 in range? {ch4_start <= 351 < ch4_end}")

# Check what's at paras 345-355
print("\nParas 345-355:")
for i in range(345, 356):
    p = paragraphs[i]
    txt = p['text'][:100] if p['text'] else '(empty)'
    print(f"  [{i}] style={p['style']:15s} '{txt}'")
