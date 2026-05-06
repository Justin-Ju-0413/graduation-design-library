"""
Fill FYP template with thesis content.
"""
import re, os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

LATEX_DIR = r"C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex"
TEMPLATE = r"C:\Users\16084\Documents\Graduation_Design_Library\11_FYP_requirement\3+1+X FYP_Report_Template_251024.docx"
OUTPUT = os.path.join(LATEX_DIR, "FYP_Thesis_Final.docx")

def read_tex(name):
    with open(os.path.join(LATEX_DIR, "chapters", name), 'r', encoding='utf-8') as f:
        return f.read()

def strip_cmd(text):
    """Strip LaTeX commands, keep readable text."""
    text = re.sub(r'(?<!\\)%.*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', '[Figure]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{table\}.*?\\end\{table\}', '[Table]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{lstlisting\}.*?\\end\{lstlisting\}', '[Code]', text, flags=re.DOTALL)
    text = re.sub(r'\\section\{([^}]*)\}', r'\n\n\1\n', text)
    text = re.sub(r'\\subsection\{([^}]*)\}', r'\n\n\1\n', text)
    text = re.sub(r'\\subsubsection\{([^}]*)\}', r'\n\1\n', text)
    text = re.sub(r'~?\\cite\{([^}]*)\}', r'', text)
    text = re.sub(r'~?\\ref\{([^}]*)\}', r'', text)
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    for cmd in ['textbf', 'texttt', 'textit', 'code', 'reg', 'emph']:
        text = re.sub(rf'\\{cmd}\{{([^}}]*)\}}', r'\1', text)
    text = re.sub(r'\$([^$]*)\$', r'\1', text)
    text = re.sub(r'\\\[.*?\\\]', '', text, flags=re.DOTALL)
    text = text.replace('\\textasciitilde{}', '~')
    text = text.replace('\\textasciicircum{}', '^')
    text = text.replace('$\\times$', 'x').replace('\\times', 'x')
    text = text.replace('\\_', '_').replace('\\#', '#')
    text = text.replace('\\&', '&')
    text = re.sub(r'\\item\s+', '\n  - ', text)
    text = re.sub(r'\\begin\{itemize\}', '', text)
    text = re.sub(r'\\end\{itemize\}', '', text)
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def add_text(doc, text, size=12, bold=False):
    """Add a paragraph with Times New Roman text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    return p

def add_section(doc, title):
    """Add a chapter heading (bold 16pt)."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    return p

def add_subsection(doc, title):
    """Add a subsection heading (bold 14pt)."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    return p

def add_paragraphs(doc, tex_text, strip_sections=True):
    """Parse cleaned text and add paragraphs, detecting section markers."""
    lines = tex_text.split('\n')
    buf = []
    for line in lines:
        line = line.strip()
        if not line:
            if buf:
                add_text(doc, ' '.join(buf), size=12)
                buf = []
            continue
        # Detect subsection markers (standalone short lines before longer text blocks)
        if len(line) < 80 and not line.startswith('-') and not line[0].isdigit():
            # Could be a subsection heading
            if re.match(r'^[A-Z].*[a-z].*$', line) and len(line.split()) <= 6:
                if buf:
                    add_text(doc, ' '.join(buf), size=12)
                    buf = []
                add_subsection(doc, line)
                continue
        buf.append(line)
    if buf:
        add_text(doc, ' '.join(buf), size=12)

print("Creating document from template...")
doc = Document(TEMPLATE)

# ================================================================
# Find key paragraph indices
# ================================================================
p_texts = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs)]

intro_i = next(i for i, t in p_texts if t == '1. Introduction')
method_i = next(i for i, t in p_texts if t.startswith('2. Methodology'))
results_i = next(i for i, t in p_texts if t.startswith('3. Results'))
disc_i = next(i for i, t in p_texts if t.startswith('4. Discussion'))
concl_i = next(i for i, t in p_texts if t.startswith('5. Conclusion'))
append_i = next(i for i, t in p_texts if t.startswith('Appendices'))
ref_i = next(i for i, t in p_texts if t.startswith('References:'))

# ================================================================
# Fill Cover Page
# ================================================================
for i, t in p_texts:
    p = doc.paragraphs[i]
    if 'Student Name:' in t:
        p.clear(); p.add_run('Student Name: JU JIAXING')
    elif 'Home University:' in t:
        p.clear(); p.add_run('Home University: City University of Hong Kong')
    elif t == 'Student ID:':
        p.clear(); p.add_run('Student ID: [Please fill in]')
    elif t == 'Supervisor:':
        p.clear(); p.add_run('Supervisor: [Supervisor Name]')
    elif t == 'Assessor:':
        p.clear(); p.add_run('Assessor: [Assessor Name]')
    elif 'Proposal Code' in t:
        p.clear(); p.add_run('2025/26-[Proposal Code]')
    elif t == 'Project Title':
        r = p.clear(); run = p.add_run('RISC-V Custom Instruction Based Lightweight CNN Accelerator: FPGA Prototype Validation')
        run.bold = True

# ================================================================
# Fill Declaration
# ================================================================
decl_i = next(i for i, t in p_texts if 'I have read the student handbook' in t)
doc.paragraphs[decl_i].clear()
doc.paragraphs[decl_i].add_run(
    'I have read the student handbook and I understand the meaning of academic dishonesty, '
    'in particular plagiarism and collusion. I declare that the work submitted for the final '
    'year project does not involve academic dishonesty. I give permission for my final year '
    'project work to be electronically scanned and if found to involve academic dishonesty, '
    'I am aware of the consequences as stated in the Student Handbook.'
)

# ================================================================
# Fill Abstract
# ================================================================
abstract_tex = read_tex('00_abstract.tex')
abstract_text = strip_cmd(abstract_tex)
# Replace the abstract paragraph (find "An abstract is an overview" paragraph)
abs_i = next(i for i, t in p_texts if 'An abstract is an overview' in t)
doc.paragraphs[abs_i].clear()
doc.paragraphs[abs_i].add_run(abstract_text)

# Remove checklist paragraphs after abstract
for i in range(abs_i+1, abs_i+12):
    if i < len(doc.paragraphs):
        doc.paragraphs[i].clear()

# ================================================================
# Clear all guide/checklist text between section headers
# ================================================================
# Clear checklist after 1. Introduction heading
for i in range(intro_i+1, method_i):
    if '1.1 Background' in doc.paragraphs[i].text or '1.2 Objectives' in doc.paragraphs[i].text:
        pass  # keep these sub-headings
    else:
        doc.paragraphs[i].clear()

# Clear between 2. Methodology and 3. Results
for i in range(method_i+1, results_i):
    doc.paragraphs[i].clear()

# Clear between 3. Results and 4. Discussion
for i in range(results_i+1, disc_i):
    doc.paragraphs[i].clear()

# Clear between 4. Discussion and 5. Conclusion
for i in range(disc_i+1, concl_i):
    doc.paragraphs[i].clear()

# Clear between 5. Conclusion and Appendices
for i in range(concl_i+1, append_i):
    doc.paragraphs[i].clear()

# Clear appendices and reference examples
for i in range(append_i, len(doc.paragraphs)):
    if doc.paragraphs[i].text.strip() in ('Appendices (Please remove this part if not applicable)',
                                           'References:'):
        pass  # keep headers
    elif 'References:' not in doc.paragraphs[i].text:
        doc.paragraphs[i].clear()

# ================================================================
# Now add content — we add as new paragraphs after the cleared areas
# Since python-docx adds at end, we need to insert in the right place.
# Best approach: add paragraphs sequentially, using the cleared template
# paragraphs as section separators.
# ================================================================

# Strategy: make all additions at the end, with section markers.
# The template's section headers remain in place as visual separators.

# Read all content
ch1_tex = read_tex('01_1_introduction.tex')
ch2_tex = read_tex('02_2_background.tex')
ch3_tex = read_tex('03_3_methodology.tex')
ch4_tex = read_tex('04_4_results.tex')
ch5_tex = read_tex('05_5_discussion.tex')
ch6_tex = read_tex('06_6_conclusion.tex')

ch1 = strip_cmd(ch1_tex)
ch2 = strip_cmd(ch2_tex)
ch3 = strip_cmd(ch3_tex)
ch4 = strip_cmd(ch4_tex)
ch5 = strip_cmd(ch5_tex)
ch6 = strip_cmd(ch6_tex)

add_section(doc, '1. Introduction')
add_subsection(doc, '1.1 Background')
add_text(doc, ch1.split('\n\n')[0] if ch1 else '', size=12)

# Background content
add_text(doc, 'RISC-V is an open standard instruction set architecture (ISA)...', size=12)
for para in ch2.split('\n\n'):
    para = para.strip()
    if not para or len(para) < 20:
        continue
    add_text(doc, para, size=12)

add_subsection(doc, '1.2 Objectives')
# Extract objectives from ch1
obj_match = re.search(r'The specific objectives are:(.*?)(?:\n\n[A-Z]|\Z)', ch1, re.DOTALL)
if obj_match:
    obj_text = obj_match.group(1).strip()
    for line in obj_text.split('\n'):
        line = line.strip().lstrip('•').strip()
        if line and len(line) > 10:
            add_text(doc, '- ' + line, size=12)

add_section(doc, '2. Methodology')
for para in ch3.split('\n\n'):
    para = para.strip()
    if not para:
        continue
    # Detect subsection
    if len(para) < 80 and para[0].isupper() and '\n' not in para and not para.startswith('[Figure]'):
        add_subsection(doc, para)
    else:
        add_text(doc, para, size=12)

add_section(doc, '3. Results')
for para in ch4.split('\n\n'):
    para = para.strip()
    if not para:
        continue
    add_text(doc, para, size=12)

add_section(doc, '4. Discussion')
for para in ch5.split('\n\n'):
    para = para.strip()
    if not para:
        continue
    add_text(doc, para, size=12)

add_section(doc, '5. Conclusion')
for para in ch6.split('\n\n'):
    para = para.strip()
    if not para:
        continue
    add_text(doc, para, size=12)

# References
add_section(doc, 'References')
bib_path = os.path.join(LATEX_DIR, 'references.bib')
with open(bib_path, 'r', encoding='utf-8') as f:
    bib = f.read()

entries = re.findall(r'@\w+\{([^,]+),\s*\n(.*?)\n\}', bib, re.DOTALL)
for n, (key, fields) in enumerate(entries, 1):
    title = re.search(r'title\s*=\s*\{([^}]*)\}', fields)
    author = re.search(r'author\s*=\s*\{([^}]*)\}', fields)
    year = re.search(r'year\s*=\s*\{([^}]*)\}', fields)
    journal = re.search(r'journal\s*=\s*\{([^}]*)\}', fields)
    booktitle = re.search(r'booktitle\s*=\s*\{([^}]*)\}', fields)
    institution = re.search(r'institution\s*=\s*\{([^}]*)\}', fields)
    school = re.search(r'school\s*=\s*\{([^}]*)\}', fields)
    vol = re.search(r'volume\s*=\s*\{([^}]*)\}', fields)
    num = re.search(r'number\s*=\s*\{([^}]*)\}', fields)
    pages = re.search(r'pages\s*=\s*\{([^}]*)\}', fields)

    a = author.group(1) if author else ''
    t = title.group(1) if title else ''
    j = journal.group(1) if journal else booktitle.group(1) if booktitle else institution.group(1) if institution else school.group(1) if school else ''
    y = year.group(1) if year else ''
    v = vol.group(1) if vol else ''
    nu = num.group(1) if num else ''
    pp = pages.group(1) if pages else ''

    ref = f'[{n}] {a}, "{t}," '
    if j:
        ref += f'{j}, '
    if v:
        ref += f'vol. {v}, '
    if nu:
        ref += f'no. {nu}, '
    if pp:
        ref += f'pp. {pp}, '
    ref += f'{y}.'
    add_text(doc, ref, size=10)

# Save
print(f"Saving...")
doc.save(OUTPUT)
print(f"Done: {OUTPUT}")
