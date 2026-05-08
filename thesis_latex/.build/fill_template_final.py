"""
Fill FYP template — V4 HYBRID: Pandoc body (tables+citations) + manual figures + FYP front matter.
"""
import re, os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

LATEX_DIR = r"C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex"
FIG_DIR = os.path.join(LATEX_DIR, "figures")
TEMPLATE = r"C:\Users\16084\Documents\Graduation_Design_Library\11_FYP_requirement\3+1+X FYP_Report_Template_251024.docx"
PANDOC_DOCX = os.path.join(LATEX_DIR, ".build", "_pandoc_full.docx")
OUTPUT = r"C:\Users\16084\Documents\Graduation_Design_Library\FYP_Thesis_Final.docx"

# Figure definitions: (filename, caption, width_inches)
ALL_FIGURES = [
    ('fig3_1_soc_architecture.png', 'E203 SoC with NICE CNN Accelerator — System Architecture', 5.5),
    ('fig3_2_instruction_format.png', 'NICE Custom Instruction Encoding Format (opcode 0x0B)', 5.0),
    ('fig3_2b_instruction_table.png', 'NICE Custom Instruction Set for the CNN Accelerator', 5.0),
    ('fig3_3_pe_microarchitecture.png', 'Processing Element (PE) Microarchitecture', 5.5),
    ('fig3_4_pe_array.png', '4×4 Systolic Processing Element Array (Output-Stationary Dataflow)', 5.5),
    ('fig3_5_packed_format.png', 'Packed INT8 Data Format for WLOAD and DLOAD Instructions', 5.0),
    ('fig3_6_build_pipeline.png', 'FPGA Bitstream Build Pipeline', 5.5),
    ('fig3_7_verification_chain.png', 'Verification Flow Chain from RTL Simulation to FPGA Board Test', 5.5),
    ('fig_fpga_board.jpg', 'Davinci Pro A7-100T FPGA Development Board', 5.0),
    ('fig4_1_ila_pc_trace.png', 'ILA Capture: hello_e203 Boot Sequence (PC Progression)', 5.5),
    ('fig4_2_ila_nice_activity.png', 'ILA Capture: NICE CNN Accelerator Instruction Execution', 5.5),
    ('fig4_3_speedup_bar.png', 'CNN Accelerator Performance Speedup (5.28× over CPU-only)', 4.5),
    ('fig4_5_utilization.png', 'FPGA Resource Utilization (Post-Implementation)', 5.0),
    ('fig4_4_resource_pie.png', 'FPGA Resource Utilization Distribution (Post-Placement)', 4.5),
    ('fig4_6_timing.png', 'Timing Closure Across FPGA Build Configurations', 5.5),
    ('fig_uart_output.png', 'UART Terminal Output: LeNet-5 MNIST Inference (10/10 Correct)', 5.5),
]


def set_tnr(para, size=12, bold=False):
    pf = para.paragraph_format
    pf.line_spacing = 1.5
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        if bold:
            run.bold = True


def add_figure(doc, filename, caption, width_inches=5.0):
    filepath = os.path.join(FIG_DIR, filename)
    if not os.path.exists(filepath):
        print(f"  SKIP missing: {filename}")
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(filepath, width=Inches(min(width_inches, 5.5)))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.add_run(f'Figure: {caption}')
    set_tnr(p_cap, size=10)


# ================================================================
# 1. BUILD FYP FRONT MATTER FROM TEMPLATE
# ================================================================
print("[1/5] Building FYP front matter...")
doc = Document(TEMPLATE)

for section in doc.sections:
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.5)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Fill cover
p_texts = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs)]
for i, t in p_texts:
    p = doc.paragraphs[i]
    if 'Student Name:' in t and 'Student ID' not in t:
        p.clear(); p.add_run('Student Name: JU JIAXING')
    elif 'Home University:' in t:
        p.clear(); p.add_run('Home University: South China University of Technology')
    elif t == 'Student ID:':
        p.clear(); p.add_run('Student ID: 202264710392')
    elif t == 'Supervisor:':
        p.clear(); p.add_run('Supervisor: [Supervisor Name]')
    elif t == 'Assessor:':
        p.clear(); p.add_run('Assessor: [Assessor Name]')
    elif 'Proposal Code' in t:
        p.clear(); p.add_run('2025/26-[Proposal Code]')
    elif t == 'Project Title':
        p.clear()
        p.add_run('RISC-V Custom Instruction Based Lightweight CNN Accelerator: FPGA Prototype Validation').bold = True

# Fill cover table
cover_table = doc.tables[0]
cover_table.rows[0].cells[0].paragraphs[0].clear()
cover_table.rows[0].cells[0].paragraphs[0].add_run(
    'RISC-V Custom Instruction Based Lightweight CNN Accelerator: FPGA Prototype Validation').bold = True
cover_table.rows[2].cells[0].paragraphs[0].clear()
cover_table.rows[2].cells[0].paragraphs[0].add_run('Student Name: JU JIAXING')
cover_table.rows[2].cells[2].paragraphs[0].clear()
cover_table.rows[2].cells[2].paragraphs[0].add_run('Student ID: 202264710392')

# Remove guide paragraphs
for i, t in p_texts:
    if 'Parts in yellow highlights' in t:
        doc.paragraphs[i].clear()
        break

# Declaration
decl_text = (
    'I have read the student handbook and I understand the meaning of academic dishonesty, '
    'in particular plagiarism and collusion. I declare that the work submitted for the final '
    'year project does not involve academic dishonesty. I give permission for my final year '
    'project work to be electronically scanned and if found to involve academic dishonesty, '
    'I am aware of the consequences as stated in the Student Handbook.'
)
for i, t in p_texts:
    if 'I have read the student handbook' in t:
        doc.paragraphs[i].clear()
        doc.paragraphs[i].add_run(decl_text)
        break
    if 'No part of this report may be reproduced' in t:
        doc.paragraphs[i].clear()
        break

# Abstract
for i, t in p_texts:
    if 'An abstract is an overview' in t:
        doc.paragraphs[i].clear()
        doc.paragraphs[i].add_run(
            'The increasing demand for energy-efficient edge AI inference has driven interest in '
            'hardware acceleration of convolutional neural networks (CNNs). This project presents '
            'the design, implementation, and FPGA validation of a lightweight CNN accelerator '
            'integrated with a RISC-V E203 processor via the NICE custom instruction interface. '
            'A 4×4 systolic processing element array supporting INT8 quantized multiply-accumulate '
            'operations was implemented in Verilog and verified through RTL simulation. The complete '
            'SoC was prototyped on a Xilinx Artix-7 FPGA, with systematic ILA-based debugging '
            'identifying and resolving four critical root causes preventing CPU boot. An end-to-end '
            'LeNet-5 inference pipeline was deployed, achieving 98.34% classification accuracy on '
            'the MNIST test set and correct classification of all 10 FPGA board test images. '
            'The accelerator achieves a 5.28× convolution speedup over CPU-only execution, '
            'demonstrating the viability of RISC-V custom instruction extensions for edge AI acceleration.'
        )
        break

# Remove ALL checklist/guide/template paragraphs
nsmap = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
body = doc.element.body
all_children = list(body)

# Find "1. Introduction" body heading
intro_body_idx = None
p_count = 0
for child in all_children:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        # Check text
        texts = [node.text or '' for node in child.iter() if node.text]
        combined = ''.join(texts).strip()
        if combined == '1. Introduction' and p_count > 70:
            intro_body_idx = child
            break
        p_count += 1

if intro_body_idx:
    removing = False
    to_remove = []
    for child in all_children:
        if child == intro_body_idx:
            removing = True
        if removing:
            to_remove.append(child)
    for child in to_remove:
        body.remove(child)
    print(f"  Removed {len(to_remove)} template body elements")

# ================================================================
# 2. LOAD PANDOC BODY (tables + text + citations)
# ================================================================
print("[2/5] Loading Pandoc-generated body (with tables)...")
pandoc_doc = Document(PANDOC_DOCX)

# ================================================================
# 3. COPY PANDOC BODY → FYP TEMPLATE, injecting figures at captions
# ================================================================
print("[3/5] Merging content with figures...")
fig_index = 0
figures_inserted = set()

for p_elem in pandoc_doc.element.body:
    tag = p_elem.tag.split('}')[-1] if '}' in p_elem.tag else p_elem.tag

    if tag == 'p':
        # Deep copy the paragraph
        new_p = etree.SubElement(doc.element.body, p_elem.tag)
        for attr_key, attr_val in p_elem.attrib.items():
            new_p.set(attr_key, attr_val)
        for child in p_elem:
            new_p.append(child)

        # Check if this paragraph is a figure caption (starts with "Figure")
        text = ''.join([node.text or '' for node in p_elem.iter() if node.text]).strip()
        if text.lower().startswith('figure') and fig_index < len(ALL_FIGURES):
            fname, caption, width = ALL_FIGURES[fig_index]
            add_figure(doc, fname, caption, width)
            fig_index += 1

    elif tag == 'tbl':
        # Deep copy the table
        new_tbl = etree.SubElement(doc.element.body, p_elem.tag)
        for attr_key, attr_val in p_elem.attrib.items():
            new_tbl.set(attr_key, attr_val)
        for child in p_elem:
            new_tbl.append(child)

    elif tag == 'sdt':
        # TOC or other structured elements — copy them
        new_sdt = etree.SubElement(doc.element.body, p_elem.tag)
        for attr_key, attr_val in p_elem.attrib.items():
            new_sdt.set(attr_key, attr_val)
        for child in p_elem:
            new_sdt.append(child)

# ================================================================
# 4. INSERT ANY MISSING FIGURES AT END
# ================================================================
print("[4/5] Adding remaining figures...")
for i in range(fig_index, len(ALL_FIGURES)):
    fname, caption, width = ALL_FIGURES[i]
    add_figure(doc, fname, caption, width)

# ================================================================
# 5. FINAL FORMATTING
# ================================================================
print("[5/5] Final formatting...")
for para in doc.paragraphs:
    if para.runs:
        for run in para.runs:
            if run.font.name is None:
                run.font.name = 'Times New Roman'
            if run.font.size is None and run.text.strip():
                run.font.size = Pt(12)
    # Set line spacing for all paragraphs
    if para.paragraph_format.line_spacing is None:
        para.paragraph_format.line_spacing = 1.5

doc.save(OUTPUT)
print(f"\nDone! Output: {OUTPUT}")
