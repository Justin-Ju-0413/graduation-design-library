"""
Fill FYP template with thesis content — clean approach.
Removes template guide text, then inserts content in correct order.
"""
import re, os, copy
from docx import Document
from docx.shared import Pt, Inches
from lxml import etree

LATEX_DIR = r"C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex"
TEMPLATE = r"C:\Users\16084\Documents\Graduation_Design_Library\11_FYP_requirement\3+1+X FYP_Report_Template_251024.docx"
OUTPUT = os.path.join(LATEX_DIR, "FYP_Thesis_Final.docx")

def read_tex(name):
    with open(os.path.join(LATEX_DIR, "chapters", name), 'r', encoding='utf-8') as f:
        return f.read()

def strip_cmd(text):
    """Strip LaTeX commands to readable text."""
    text = re.sub(r'(?<!\\)%.*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', '', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{table\}.*?\\end\{table\}', '', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{lstlisting\}.*?\\end\{lstlisting\}', '', text, flags=re.DOTALL)
    text = re.sub(r'\\section\{([^}]*)\}', r'\n\n\1\n', text)
    text = re.sub(r'\\subsection\{([^}]*)\}', r'\n\n\1\n', text)
    text = re.sub(r'\\subsubsection\{([^}]*)\}', r'\n\1\n', text)
    text = re.sub(r'~?\\cite\{([^}]*)\}', r'', text)
    text = re.sub(r'~?\\ref\{([^}]*)\}', r'', text)
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    for cmd in ['textbf', 'texttt', 'textit', 'code', 'reg', 'emph']:
        text = re.sub(rf'\\{cmd}\{{([^}}]*)\}}', r'\1', text)
    text = re.sub(r'\$([^$]*)\$', r'\1', text)
    text = re.sub(r'\\\[.*?\\\]', '', text, flags=re.DOTALL)
    text = text.replace('\\textasciitilde{}', '~')
    text = text.replace('\\textasciicircum{}', '^')
    text = text.replace('$\\times$', 'x').replace('\\times', 'x')
    text = text.replace('\\_', '_').replace('\\#', '#')
    text = text.replace('\\&', '&')
    text = re.sub(r'\\item\s+', '\n  - ', text)
    text = re.sub(r'\\begin\{itemize\}', '', text)
    text = re.sub(r'\\end\{itemize\}', '', text)
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def add_section(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    return p

def add_subsection(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    return p

def add_body(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    return p


print("Loading template...")
doc = Document(TEMPLATE)

# Identify paragraphs to keep (cover page + declaration) vs remove (guide content)
# Find where the actual template structure ends and our content should begin
# We'll keep everything up through Abstract, then rebuild from TOC onward

# Find key indices
p_meta = [(i, p.text.strip()[:100]) for i, p in enumerate(doc.paragraphs)]

# Find "1. Introduction" heading in template TOC (the first one)
intro_toc_idx = None
for i, t in p_meta:
    if t.startswith('1. Introduction'):
        intro_toc_idx = i
        break

# Find "List of Figures" start
lof_idx = None
for i, t in p_meta:
    if t.strip() == 'List of Figures':
        lof_idx = i
        break

# Strategy: Remove everything from "1. Introduction" in TOC to end
# Then rebuild: TOC + all content

# First, handle the body element to remove paragraphs after intro_toc_idx
body = doc.element.body
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

if intro_toc_idx is not None:
    # Remove paragraphs from intro_toc_idx to end
    all_p_elements = body.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    # Also remove tables
    all_tbl_elements = body.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl')

    # Find the element index in the body children
    children = list(body)
    remove_from = None
    count = 0
    for child in children:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            if count == intro_toc_idx:
                remove_from = child
                break
            count += 1

    if remove_from is not None:
        # Remove all siblings from remove_from to end
        parent = body
        removing = False
        to_remove = []
        for child in parent:
            if child == remove_from:
                removing = True
            if removing:
                to_remove.append(child)
        for child in to_remove:
            parent.remove(child)

print("Cleaned template guide content.")

# Now rebuild content from TOC onwards

# ---- Table of Contents ----
add_section(doc, 'Contents')
toc_items = [
    ('Abstract', 'i'),
    ('Acknowledgements', 'ii'),
    ('Contents', 'iii'),
    ('List of Figures', 'iv'),
    ('List of Tables', 'v'),
    ('1. Introduction', '1'),
    ('  1.1 Background', '1'),
    ('  1.2 Objectives', '2'),
    ('2. Methodology', '3'),
    ('3. Results', '14'),
    ('4. Discussion', '32'),
    ('5. Conclusion', '36'),
    ('References', '38'),
]
for item, page in toc_items:
    add_body(doc, f'{item} {"." * (50 - len(item))} {page}', size=12)

add_body(doc, '', size=12)

# ---- List of Figures ----
add_section(doc, 'List of Figures')
figures = [
    ('Figure 3.1', 'E203 SoC with CNN NICE Accelerator'),
    ('Figure 3.2', 'NICE Custom Instruction Encoding (opcode 0x0B)'),
    ('Figure 3.3', 'Processing Element (PE) Microarchitecture'),
    ('Figure 3.4', '4x4 PE Array Organization (Output Stationary)'),
    ('Figure 3.5', 'Packed INT8 Data Format (WLOAD / DLOAD)'),
    ('Figure 3.6', 'FPGA Bitstream Build Pipeline'),
    ('Figure 4.1', 'ILA Capture: hello_e203 Boot Sequence (PC Progression)'),
    ('Figure 4.2', 'ILA Capture: NICE CNN Accelerator Instruction Execution'),
    ('Figure 4.3', 'CNN Accelerator Performance Speedup'),
    ('Figure 4.4', 'FPGA Resource Utilization'),
    ('Figure 4.5', 'Timing Closure Across FPGA Build Configurations'),
    ('Figure 4.6', 'FPGA Board: Davinci Pro A7-100T'),
    ('Figure 4.7', 'UART Output: LeNet-5 MNIST Inference Demo'),
]
for num, desc in figures:
    add_body(doc, f'{num}  {desc}', size=12)

add_body(doc, '', size=12)

# ---- List of Tables ----
add_section(doc, 'List of Tables')
tables = [
    ('Table 2.1', 'E203 Processor Core Parameters'),
    ('Table 2.2', 'E203 Memory Map (Davinci A7-100T Configuration)'),
    ('Table 3.1', 'SoC Address Space Map'),
    ('Table 3.2', 'NICE Request and Response Channel Signals'),
    ('Table 3.3', 'NICE Custom Instruction Set for the CNN Accelerator'),
    ('Table 4.1', 'Summary of Build Configuration Fixes for FPGA Compatibility'),
    ('Table 4.2', 'FPGA Build Modes and Their Descriptions'),
    ('Table 4.3a', 'Timing Closure Results Across FPGA Build Configurations'),
    ('Table 4.3b', 'hello_e203 ILA Probe Values'),
    ('Table 4.4', 'NICE Accelerator ILA Probe Values'),
    ('Table 4.5', 'FPGA Resource Utilization for the Complete SoC'),
    ('Table 5.1', 'Project Objective Status Summary'),
]
for num, desc in tables:
    add_body(doc, f'{num}  {desc}', size=12)

# ---- Page break before content ----
doc.add_page_break()

# ================================================================
# CHAPTER CONTENT
# ================================================================

# Read all chapters
ch1 = strip_cmd(read_tex('01_1_introduction.tex'))
ch2 = strip_cmd(read_tex('02_2_background.tex'))
ch3 = strip_cmd(read_tex('03_3_methodology.tex'))
ch4 = strip_cmd(read_tex('04_4_results.tex'))
ch5 = strip_cmd(read_tex('05_5_discussion.tex'))
ch6 = strip_cmd(read_tex('06_6_conclusion.tex'))

# ---- Chapter 1: Introduction ----
add_section(doc, '1. Introduction')
add_subsection(doc, '1.1 Background')
# ch1 content
for para in ch1.split('\n\n'):
    para = para.strip()
    if not para or len(para) < 30:
        continue
    # Skip section headers (already handled)
    if para in ('Background', 'RISC-V and Custom Instruction Extensions',
                'Problem Statement', 'Project Objectives', 'Thesis Structure'):
        continue
    add_body(doc, para, size=12)

# Merge ch2 (Background) into Introduction as additional background
add_body(doc, '', size=6)
for para in ch2.split('\n\n'):
    para = para.strip()
    if not para or len(para) < 30:
        continue
    add_body(doc, para, size=12)

add_subsection(doc, '1.2 Objectives')
# Extract objectives
obj_match = re.search(r'The specific objectives are:(.*?)(?:Thesis Structure|\Z)', ch1, re.DOTALL)
if obj_match:
    obj_text = obj_match.group(1).strip()
    for line in obj_text.split('\n'):
        line = line.strip().lstrip('•').lstrip('-').strip()
        if line and len(line) > 10:
            add_body(doc, '- ' + line, size=12)

# ---- Chapter 2: Methodology ----
add_section(doc, '2. Methodology')
for para in ch3.split('\n\n'):
    para = para.strip()
    if not para or len(para) < 20:
        continue
    # Skip known section headers that were stripped
    add_body(doc, para, size=12)

# ---- Chapter 3: Results ----
add_section(doc, '3. Results')
for para in ch4.split('\n\n'):
    para = para.strip()
    if not para or len(para) < 10:
        continue
    add_body(doc, para, size=12)

# ---- Chapter 4: Discussion ----
add_section(doc, '4. Discussion')
for para in ch5.split('\n\n'):
    para = para.strip()
    if not para or len(para) < 10:
        continue
    add_body(doc, para, size=12)

# ---- Chapter 5: Conclusion ----
add_section(doc, '5. Conclusion')
for para in ch6.split('\n\n'):
    para = para.strip()
    if not para or len(para) < 10:
        continue
    add_body(doc, para, size=12)

# ---- References ----
add_section(doc, 'References')

bib_path = os.path.join(LATEX_DIR, 'references.bib')
with open(bib_path, 'r', encoding='utf-8') as f:
    bib = f.read()

entries = re.findall(r'@\w+\{([^,]+),\s*\n(.*?)\n\}', bib, re.DOTALL)
for n, (key, fields) in enumerate(entries, 1):
    title = re.search(r'title\s*=\s*\{([^}]*)\}', fields)
    author = re.search(r'author\s*=\s*\{([^}]*)\}', fields)
    year = re.search(r'year\s*=\s*\{([^}]*)\}', fields)
    journal = re.search(r'journal\s*=\s*\{([^}]*)\}', fields)
    booktitle = re.search(r'booktitle\s*=\s*\{([^}]*)\}', fields)
    institution = re.search(r'institution\s*=\s*\{([^}]*)\}', fields)
    school = re.search(r'school\s*=\s*\{([^}]*)\}', fields)
    publisher = re.search(r'publisher\s*=\s*\{([^}]*)\}', fields)
    vol = re.search(r'volume\s*=\s*\{([^}]*)\}', fields)
    num = re.search(r'number\s*=\s*\{([^}]*)\}', fields)
    pages = re.search(r'pages\s*=\s*\{([^}]*)\}', fields)
    note = re.search(r'note\s*=\s*\{([^}]*)\}', fields)

    a = author.group(1) if author else ''
    t = title.group(1) if title else ''
    y = year.group(1) if year else ''
    pub = journal.group(1) if journal else booktitle.group(1) if booktitle else institution.group(1) if institution else school.group(1) if school else publisher.group(1) if publisher else ''
    v = vol.group(1) if vol else ''
    nu = num.group(1) if num else ''
    pp = pages.group(1) if pages else ''
    nt = note.group(1) if note else ''

    ref = f'[{n}] {a}, "{t}," '
    if pub:
        ref += f'{pub}, '
    if v:
        ref += f'vol. {v}, '
    if nu:
        ref += f'no. {nu}, '
    if pp:
        ref += f'pp. {pp}, '
    ref += f'{y}.'
    if nt:
        ref += f' {nt}.'
    add_body(doc, ref, size=10)

# Save
print(f"Saving...")
doc.save(OUTPUT)
print(f"Done: {OUTPUT}")
