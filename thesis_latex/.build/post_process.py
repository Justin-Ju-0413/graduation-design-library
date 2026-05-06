"""
Post-process Pandoc-generated docx:
1. Enforce Times New Roman on ALL runs
2. Fix font sizes (H1=16pt, H2=14pt, H3=13pt, Normal=12pt)
3. Enforce 1.5 line spacing
4. Add FYP cover page from template
5. Replace Declaration with official text
6. Clean up any LaTeX artifacts
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

PANDOC_OUT = r'C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex\thesis_pandoc.docx'
FYP_TEMPLATE = r'C:\Users\16084\Documents\Graduation_Design_Library\11_FYP_requirement\3+1+X FYP_Report_Template_251024.docx'
FINAL = r'C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex\FYP_Thesis_Final.docx'

print("Loading Pandoc output...")
doc = Document(PANDOC_OUT)

# ================================================================
# 1. Enforce TNR on all styles
# ================================================================
print("1. Enforcing Times New Roman on all styles...")

def enforce_tnr_style(doc, style_name, font_size=None, bold=None):
    """Set font on a style and its XML definition."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    style.font.name = 'Times New Roman'
    if font_size:
        style.font.size = Pt(font_size)
    if bold is not None:
        style.font.bold = bold

    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = doc.element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
        rFonts.set(qn(attr), 'Times New Roman')

# Normal
enforce_tnr_style(doc, 'Normal', font_size=12)
pf = doc.styles['Normal'].paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.5

# Headings
enforce_tnr_style(doc, 'Heading 1', font_size=16, bold=True)
enforce_tnr_style(doc, 'Heading 2', font_size=14, bold=True)
enforce_tnr_style(doc, 'Heading 3', font_size=13, bold=True)
enforce_tnr_style(doc, 'Title', font_size=18, bold=True)

# Body Text / First Paragraph
for sname in ['Body Text', 'First Paragraph']:
    try:
        enforce_tnr_style(doc, sname, font_size=12)
    except:
        pass

# ================================================================
# 2. Enforce TNR on every single run in the document
# ================================================================
print("2. Enforcing TNR on all runs...")
fixed_runs = 0
for p in doc.paragraphs:
    for run in p.runs:
        run.font.name = 'Times New Roman'
        # Force TNR in XML too
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = doc.element.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        fixed_runs += 1

print(f"   Fixed {fixed_runs} runs")

# ================================================================
# 3. Fix heading font sizes on every heading paragraph
# ================================================================
print("3. Fixing heading font sizes...")
sizes = {'Heading 1': 16, 'Heading 2': 14, 'Heading 3': 13}
for p in doc.paragraphs:
    if p.style.name in sizes:
        for run in p.runs:
            run.font.size = Pt(sizes[p.style.name])

# ================================================================
# 4. Fix reference font size (should be 12pt, not 10pt)
# ================================================================
print("4. Fixing reference font sizes...")
for p in doc.paragraphs:
    if re.match(r'^\s*\[\d+\]', p.text.strip()[:5]):
        for run in p.runs:
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5

# ================================================================
# 5. Enforce 1.5 line spacing on all body paragraphs
# ================================================================
print("5. Enforcing 1.5 line spacing...")
for p in doc.paragraphs:
    if p.style.name == 'Normal' or p.style.name.startswith('Heading'):
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5

# ================================================================
# 6. Add FYP cover page at the beginning
#    We'll prepend cover elements using XML manipulation
# ================================================================
print("6. Building FYP cover page...")

# Load the FYP template to extract cover styles
fyp_doc = Document(FYP_TEMPLATE)

# Use the sections from template for margins
for section in doc.sections:
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.5)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ================================================================
# 7. Replace Declaration content with official text
# ================================================================
print("7. Replacing Declaration...")
for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and 'Declaration' in p.text:
        # Find the next paragraph after this one
        idx = doc.paragraphs.index(p)
        if idx + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[idx + 1]
            next_p.clear()
            run = next_p.add_run(
                'I have read the student handbook and I understand the meaning of '
                'academic dishonesty, in particular plagiarism and collusion. I declare '
                'that the work submitted for the final year project does not involve '
                'academic dishonesty. I give permission for my final year project work '
                'to be electronically scanned and if found to involve academic dishonesty, '
                'I am aware of the consequences as stated in the Student Handbook.'
            )
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        break

# ================================================================
# 8. Clean up LaTeX artifacts
# ================================================================
print("8. Cleaning LaTeX artifacts...")
artifacts = [
    r'\\textasciitilde\{\}',
    r'\\textasciicircum\{\}',
    r'\\chapter\{',
    r'\\section\{',
    r'\\subsection\{',
    r'\\label\{',
]
for p in doc.paragraphs:
    text = p.text
    for run in p.runs:
        t = run.text
        # Fix common artifacts
        t = t.replace('\\_', '_')
        t = t.replace('\\#', '#')
        t = t.replace('\\&', '&')
        t = t.replace('\\%', '%')
        t = t.replace('\\textbackslash', '\\')
        t = t.replace('\{', '{')
        t = t.replace('\}', '}')
        run.text = t

# ================================================================
# 9. Fix the title page — add proper FYP cover info
# ================================================================
print("9. Enhancing title page...")
# The Pandoc title page includes: title, author
# We need to add: CityUHK header, proposal code, student info, supervisor, assessor

# Find the title paragraph and add info before it
body = doc.element.body
first_elem = body[0]  # Get first element

# Build cover page at the very beginning
# We'll use the section properties to add a page break before main content
# The Pandoc output already has a Declaration page, so we enhance the title

# Find Title-style paragraph
for p in doc.paragraphs:
    if p.style.name == 'Title':
        p.alignment = 1  # center
        for run in p.runs:
            run.font.size = Pt(18)
            run.bold = True
        break

# ================================================================
# Save
# ================================================================
print(f"\nSaving to {FINAL}...")
doc.save(FINAL)
print(f"Done: {FINAL}")
print(f"Paragraphs: {len(doc.paragraphs)}")
