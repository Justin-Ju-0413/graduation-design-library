"""
Post-process Pandoc-generated DOCX for thesis submission.
- Enforces Times New Roman 12pt on all text
- Ensures Heading 1=16pt, Heading 2=14pt, Heading 3=13pt (all bold)
- Sets 1.5 line spacing
- Fixes references to 12pt
- Replaces Declaration and Abstract with official text
- Cleans LaTeX artifacts
- Sets margins: top/bottom 1.5in, left/right 1.0in
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BUILD_DIR = SCRIPT_DIR  # .build/
LATEX_DIR = os.path.dirname(SCRIPT_DIR)  # thesis_latex/

SRC = os.path.join(BUILD_DIR, "_pandoc_raw.docx")
DST = os.path.join(LATEX_DIR, "FYP_FINAL.docx")

if not os.path.exists(SRC):
    print(f"ERROR: Pandoc output not found at {SRC}")
    print("Run Pandoc first: pandoc main.tex -o .build/_pandoc_raw.docx ...")
    exit(1)

print(f"Loading {SRC}...")
doc = Document(SRC)

# ---- 1. Configure built-in styles ----
print("  Configuring styles...")
for name, size, bold in [('Normal', 12, None), ('Body Text', 12, None),
                          ('First Paragraph', 12, None),
                          ('Heading 1', 16, True), ('Heading 2', 14, True),
                          ('Heading 3', 13, True)]:
    try:
        s = doc.styles[name]
        s.font.name = 'Times New Roman'
        if size: s.font.size = Pt(size)
        if bold is not None: s.font.bold = bold
        rpr = s.element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = doc.element.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rf)
        for a in ['w:ascii', 'w:hAnsi', 'w:cs']:
            rf.set(qn(a), 'Times New Roman')
    except Exception:
        pass

from docx.enum.text import WD_ALIGN_PARAGRAPH
for sname in ['Normal', 'Body Text', 'First Paragraph']:
    try:
        doc.styles[sname].paragraph_format.line_spacing = 1.5
        doc.styles[sname].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except Exception:
        pass

# ---- 2. Enforce fonts/sizes on every run ----
print("  Enforcing fonts on all runs...")
for p in doc.paragraphs:
    try:
        runs = p.runs
    except AttributeError:
        continue
    try:
        style = p.style.name
    except AttributeError:
        continue

    for r in runs:
        r.font.name = 'Times New Roman'
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = doc.element.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rf)
        for a in ['w:ascii', 'w:hAnsi', 'w:cs']:
            rf.set(qn(a), 'Times New Roman')

    if style == 'Heading 1':
        for r in runs:
            r.font.size = Pt(16)
    elif style == 'Heading 2':
        for r in runs:
            r.font.size = Pt(14)
    elif style == 'Heading 3':
        for r in runs:
            r.font.size = Pt(13)

    if 'Heading' in style or style in ('Normal', 'Body Text', 'First Paragraph'):
        p.paragraph_format.line_spacing = 1.5
    # Justify all body text (skip headings, references)
    if style in ('Normal', 'Body Text', 'First Paragraph') \
       and len(p.text.strip()) > 30 \
       and not re.match(r'^\[\d+\]', p.text.strip()[:10]):
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # References: force 12pt, left-aligned
    if re.match(r'^\[\d+\]', p.text.strip()[:10]):
        for r in runs:
            r.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ---- 3. Fix Declaration text ----
print("  Fixing Declaration...")
for i, p in enumerate(doc.paragraphs):
    try:
        style = p.style.name
    except AttributeError:
        continue
    if style == 'Heading 1' and 'Declaration' in p.text and i + 1 < len(doc.paragraphs):
        np = doc.paragraphs[i + 1]
        np.clear()
        r = np.add_run(
            "I have read the student handbook and I understand the meaning of "
            "academic dishonesty, in particular plagiarism and collusion. I declare "
            "that the work submitted for the final year project does not involve "
            "academic dishonesty. I give permission for my final year project work "
            "to be electronically scanned and if found to involve academic dishonesty, "
            "I am aware of the consequences as stated in the Student Handbook."
        )
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        break

# ---- 4. Fix Abstract (ensure full text is present) ----
print("  Checking Abstract...")
for i, p in enumerate(doc.paragraphs):
    try:
        style = p.style.name
    except AttributeError:
        continue
    if style == 'Heading 1' and 'Abstract' in p.text and i + 1 < len(doc.paragraphs):
        ap = doc.paragraphs[i + 1]
        if len(ap.text.strip()) < 200:
            print("    Abstract too short, reinserting...")
            ap.clear()
            r = ap.add_run(
                "The increasing demand for energy-efficient edge AI inference has "
                "driven interest in hardware acceleration of convolutional neural "
                "networks (CNNs). RISC-V's open instruction set architecture provides "
                "a unique opportunity to extend a standard processor with custom "
                "instructions for domain-specific acceleration. This project presents "
                "the design, implementation, and FPGA-based prototype validation of a "
                "lightweight CNN accelerator integrated with a RISC-V E203 processor "
                "through the Nuclei Instruction Co-extension (NICE) interface. The "
                "accelerator implements a 4x4 systolic PE array supporting INT8 "
                "quantized convolution, with six custom NICE instructions. The design "
                "was validated through RTL simulation, full-SoC simulation, and FPGA "
                "prototype bring-up on the Davinci Pro A7-100T board. A bare-metal "
                "program was validated on the FPGA, ILA debugging resolved four "
                "critical root causes, NICE test programs confirmed correct instruction "
                "execution, and an end-to-end LeNet-5 inference pipeline achieved 70 "
                "percent accuracy on MNIST. This project demonstrates a complete FPGA "
                "bring-up pipeline for a RISC-V-based CNN accelerator."
            )
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
        break

# ---- 5. Clean LaTeX artifacts ----
print("  Cleaning LaTeX artifacts...")
for p in doc.paragraphs:
    try:
        runs = p.runs
    except AttributeError:
        continue
    for r in runs:
        t = r.text
        for o, n in [(chr(92)+'_', '_'), (chr(92)+'#', '#'), (chr(92)+'&', '&'),
                      (chr(92)+'%', '%'), (chr(92)+'{', '{'), (chr(92)+'}', '}')]:
            t = t.replace(o, n)
        r.text = t

# ---- 6. Page margins ----
print("  Setting page margins...")
for s in doc.sections:
    s.top_margin = Inches(1.5)
    s.bottom_margin = Inches(1.5)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)

# ---- Save ----
print(f"  Saving {DST}...")
doc.save(DST)

# ---- Quick verify ----
doc2 = Document(DST)
wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
imgs = len(doc2.element.findall('.//{%s}inline' % wp))

bad_font = 0
for p in doc2.paragraphs:
    try:
        for r in p.runs:
            if r.font.name and r.font.name != 'Times New Roman':
                bad_font += 1
    except Exception:
        continue

refs = sum(1 for p in doc2.paragraphs if re.match(r'^\[\d+\]', p.text.strip()[:10]))
size_kb = os.path.getsize(DST) // 1024

print(f"  Done: {size_kb}KB, {len(doc2.paragraphs)} paras, "
      f"{imgs} images, {refs} refs, {bad_font} non-TNR")
