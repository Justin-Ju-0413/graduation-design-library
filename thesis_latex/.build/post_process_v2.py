"""
Post-process Pandoc docx:
1. Insert FYP cover page at beginning
2. Enforce TNR on ALL runs, proper sizes, 1.5 spacing
3. Replace Declaration, clean artifacts
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

PANDOC_OUT = r'C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex\thesis_pandoc.docx'
FINAL = r'C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex\FYP_Thesis_Final.docx'

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

print("Loading Pandoc output...")
doc = Document(PANDOC_OUT)


def safe_runs(p):
    try: return safe_runs(p)
    except AttributeError: return []

def safe_style(p):
    try: return safe_style(p)
    except AttributeError: return None

# ================================================================
# PART A: Insert FYP Cover Page at the very beginning
# ================================================================
print("A. Building FYP cover page...")

def make_cover_para(doc, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6)):
    """Create a cover paragraph element."""
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    if bold: run.bold = True
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(0)
    return p

# We'll collect cover paragraphs and insert them at the beginning
# First, save the existing content, clear the doc, add cover, then re-add

# Actually, simpler: insert before first paragraph using XML
body = doc.element.body

# Build cover paragraphs as XML
def make_p_element(text, font_size=12, bold=False, align='center', font_name='Times New Roman'):
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Paragraph properties
    pPr = etree.SubElement(etree.Element(f'{{{ns}}}pPr'), f'{{{ns}}}jc')
    pPr.set(f'{{{ns}}}val', align)

    pPr_elem = etree.Element(f'{{{ns}}}pPr')
    jc = etree.SubElement(pPr_elem, f'{{{ns}}}jc')
    jc.set(f'{{{ns}}}val', align)

    # Run properties
    rPr = etree.Element(f'{{{ns}}}rPr')
    rFonts = etree.SubElement(rPr, f'{{{ns}}}rFonts')
    rFonts.set(f'{{{ns}}}ascii', font_name)
    rFonts.set(f'{{{ns}}}hAnsi', font_name)
    rFonts.set(f'{{{ns}}}cs', font_name)
    sz = etree.SubElement(rPr, f'{{{ns}}}sz')
    sz.set(f'{{{ns}}}val', str(font_size * 2))  # half-points
    szCs = etree.SubElement(rPr, f'{{{ns}}}szCs')
    szCs.set(f'{{{ns}}}val', str(font_size * 2))
    if bold:
        b = etree.SubElement(rPr, f'{{{ns}}}b')

    # Text run
    r = etree.Element(f'{{{ns}}}r')
    r.append(rPr)
    t = etree.SubElement(r, f'{{{ns}}}t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # Paragraph
    p = etree.Element(f'{{{ns}}}p')
    p.append(pPr_elem)
    p.append(r)
    return p

def make_empty_p():
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p = etree.Element(f'{{{ns}}}p')
    pPr = etree.SubElement(p, f'{{{ns}}}pPr')
    jc = etree.SubElement(pPr, f'{{{ns}}}jc')
    jc.set(f'{{{ns}}}val', 'center')
    return p

def make_br_p():
    """Page break paragraph."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p = etree.Element(f'{{{ns}}}p')
    pPr = etree.SubElement(p, f'{{{ns}}}pPr')
    sectPr = etree.SubElement(pPr, f'{{{ns}}}sectPr')
    pgSz = etree.SubElement(sectPr, f'{{{ns}}}pgSz')
    pgSz.set(f'{{{ns}}}w', '11906')  # A4 width in twips
    pgSz.set(f'{{{ns}}}h', '16838')
    # Set margins with section properties
    pm = etree.SubElement(sectPr, f'{{{ns}}}pgMar')
    pm.set(f'{{{ns}}}top', '2743')     # 1.5in in twips (1440*1.5)
    pm.set(f'{{{ns}}}bottom', '2743')
    pm.set(f'{{{ns}}}left', '1440')    # 1.0in
    pm.set(f'{{{ns}}}right', '1440')
    return p

# Build cover paragraphs
cover_elements = []

# CityUHK header
cover_elements.append(make_p_element('CityUHK Qingdao Research Institute', 13, False))
cover_elements.append(make_empty_p())
cover_elements.append(make_p_element('FINAL YEAR PROJECT REPORT', 16, True))
cover_elements.append(make_empty_p())
cover_elements.append(make_p_element('2025/26-[Proposal Code]', 12, False))
cover_elements.append(make_empty_p())
cover_elements.append(make_p_element('RISC-V Custom Instruction Based', 16, True))
cover_elements.append(make_p_element('Lightweight CNN Accelerator:', 16, True))
cover_elements.append(make_p_element('FPGA Prototype Validation', 16, True))
cover_elements.append(make_empty_p())

# Student info (left aligned)
def make_info_p(text):
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p = etree.Element(f'{{{ns}}}p')
    pPr = etree.SubElement(p, f'{{{ns}}}pPr')
    jc = etree.SubElement(pPr, f'{{{ns}}}jc')
    jc.set(f'{{{ns}}}val', 'left')
    pPr2 = etree.SubElement(pPr, f'{{{ns}}}ind')
    pPr2.set(f'{{{ns}}}left', '4320')  # 3 inch indent
    r = etree.Element(f'{{{ns}}}r')
    rPr = etree.SubElement(r, f'{{{ns}}}rPr')
    rFonts = etree.SubElement(rPr, f'{{{ns}}}rFonts')
    rFonts.set(f'{{{ns}}}ascii', 'Times New Roman')
    rFonts.set(f'{{{ns}}}hAnsi', 'Times New Roman')
    sz = etree.SubElement(rPr, f'{{{ns}}}sz')
    sz.set(f'{{{ns}}}val', '24')  # 12pt
    t = etree.SubElement(r, f'{{{ns}}}t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    p.append(r)
    return p

info_items = [
    'Student Name:  JU JIAXING',
    'Home University:  City University of Hong Kong',
    'Student ID:  [Please fill in]',
    'Supervisor:  [Supervisor Name]',
    'Assessor:  [Assessor Name]',
    'Major:  Computer Engineering',
    'Programme:  BEng Computer Engineering',
]
for item in info_items:
    cover_elements.append(make_info_p(item))

cover_elements.append(make_empty_p())
cover_elements.append(make_p_element('Department of Computer Science', 12, False))
cover_elements.append(make_p_element('City University of Hong Kong', 12, False))
cover_elements.append(make_empty_p())
cover_elements.append(make_p_element('May 2026', 12, False))

# Page break after cover
cover_elements.append(make_br_p())

# Insert all cover elements at the beginning of body
first_child = body[0] if len(body) > 0 else None
for elem in reversed(cover_elements):
    if first_child is not None:
        body.insert(0, elem)
    else:
        body.append(elem)

print("   Cover page inserted.")

# ================================================================
# PART B: Enforce TNR on all styles
# ================================================================
print("B. Enforcing Times New Roman on all styles...")

def set_style_font(doc, name, size=None, bold=None):
    try:
        s = doc.styles[name]
    except KeyError:
        return
    s.font.name = 'Times New Roman'
    if size: s.font.size = Pt(size)
    if bold is not None: s.font.bold = bold
    rpr = s.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = doc.element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    for a in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
        rFonts.set(qn(a), 'Times New Roman')

set_style_font(doc, 'Normal', size=12)
doc.styles['Normal'].paragraph_format.line_spacing = 1.5
set_style_font(doc, 'Heading 1', size=16, bold=True)
set_style_font(doc, 'Heading 2', size=14, bold=True)
set_style_font(doc, 'Heading 3', size=13, bold=True)
set_style_font(doc, 'Title', size=18, bold=True)

# ================================================================
# PART C: Enforce on all runs
# ================================================================
print("C. Enforcing TNR on all runs...")
count = 0
for p in doc.paragraphs:
    try:
        runs = p.runs
    except AttributeError:
        continue
    for r in runs:
        r.font.name = 'Times New Roman'
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = doc.element.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rf)
        for a in ['w:ascii', 'w:hAnsi', 'w:cs']:
            rf.set(qn(a), 'Times New Roman')
        count += 1

    # Fix heading sizes
    if safe_style(p) == 'Heading 1':
        for r in runs: r.font.size = Pt(16)
    elif safe_style(p) == 'Heading 2':
        for r in runs: r.font.size = Pt(14)
    elif safe_style(p) == 'Heading 3':
        for r in runs: r.font.size = Pt(13)

    # Enforce line spacing
    if 'Heading' in safe_style(p) or safe_style(p) == 'Normal':
        p.paragraph_format.line_spacing = 1.5

print(f"   {count} runs processed")

# ================================================================
# PART D: Fix references (12pt, 1.5 spacing)
# ================================================================
print("D. Fixing references...")
ref_count = 0
for p in doc.paragraphs:
    try:
        runs = p.runs
    except AttributeError:
        continue
    if re.match(r'^\[\d+\]', p.text.strip()[:10]):
        for r in runs:
            r.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        ref_count += 1
print(f"   {ref_count} references fixed")

# ================================================================
# PART E: Replace Declaration content
# ================================================================
print("E. Replacing Declaration...")
for i, p in enumerate(doc.paragraphs):
    if safe_style(p) == 'Heading 1' and 'Declaration' in p.text:
        if i + 1 < len(doc.paragraphs):
            np = doc.paragraphs[i + 1]
            np.clear()
            r = np.add_run(
                'I have read the student handbook and I understand the meaning of '
                'academic dishonesty, in particular plagiarism and collusion. I declare '
                'that the work submitted for the final year project does not involve '
                'academic dishonesty. I give permission for my final year project work '
                'to be electronically scanned and if found to involve academic dishonesty, '
                'I am aware of the consequences as stated in the Student Handbook.'
            )
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
        break

# ================================================================
# PART F: Fix Abstract content if needed
# ================================================================
print("F. Verifying Abstract...")
for i, p in enumerate(doc.paragraphs):
    if safe_style(p) == 'Heading 1' and 'Abstract' in p.text:
        if i + 1 < len(doc.paragraphs):
            ap = doc.paragraphs[i + 1]
            if len(ap.text) < 200:
                print("   Abstract too short, reinserting...")
                ap.clear()
                r = ap.add_run(
                    'The increasing demand for energy-efficient edge AI inference has driven '
                    'interest in hardware acceleration of convolutional neural networks (CNNs). '
                    'RISC-V\'s open instruction set architecture provides a unique opportunity to '
                    'extend a standard processor with custom instructions for domain-specific '
                    'acceleration. This project presents the design, implementation, and FPGA-based '
                    'prototype validation of a lightweight CNN accelerator integrated with a RISC-V '
                    'E203 (Hummingbird v2) processor through the Nuclei Instruction Co-extension '
                    '(NICE) interface. The accelerator implements a 4x4 systolic processing element '
                    '(PE) array supporting INT8 quantized convolution operations. Six custom NICE '
                    'instructions provide the software interface. The design was validated through '
                    'RTL simulation, full-SoC simulation, and FPGA prototype bring-up on the Davinci '
                    'Pro A7-100T development board. A bare-metal program was validated on the FPGA '
                    'board, and ILA debugging resolved four critical root causes. Custom NICE test '
                    'programs confirmed correct instruction execution, and an end-to-end LeNet-5 '
                    'inference pipeline achieved 70% classification accuracy on MNIST. This project '
                    'demonstrates a complete FPGA bring-up pipeline for a RISC-V-based CNN '
                    'accelerator, establishing a reproducible evidence chain from RTL simulation '
                    'through board-level validation.'
                )
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
        break

# ================================================================
# PART G: Clean LaTeX artifacts from text
# ================================================================
print("G. Cleaning LaTeX artifacts...")
for p in doc.paragraphs:
    try:
        runs = p.runs
    except AttributeError:
        continue
    for r in runs:
        t = r.text
        t = t.replace('\\_', '_')
        t = t.replace('\\#', '#')
        t = t.replace('\\&', '&')
        t = t.replace('\\%', '%')
        t = t.replace('\\{', '{')
        t = t.replace('\\}', '}')
        t = t.replace('\\textbackslash{}', '\\')
        r.text = t

# ================================================================
# PART H: Page margins (ensure all sections)
# ================================================================
print("H. Setting page margins...")
for section in doc.sections:
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.5)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ================================================================
# Save
# ================================================================
print(f"\nSaving {FINAL}...")
doc.save(FINAL)

# Quick stats
import os
size_mb = os.path.getsize(FINAL) / (1024*1024)
paras = len(doc.paragraphs)
print(f"Done: {FINAL}")
print(f"Size: {size_mb:.1f} MB, Paragraphs: {paras}")
