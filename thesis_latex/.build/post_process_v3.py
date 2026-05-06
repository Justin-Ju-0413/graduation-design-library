"""
Post-process Pandoc docx - v3 (clean, no raw XML paragraphs)
1. Enhance Pandoc title page with FYP info
2. Enforce TNR, sizes, spacing
3. Fix references
4. Replace Declaration, fix Abstract
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

PANDOC_OUT = r'C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex\thesis_pandoc.docx'
FINAL = r'C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex\FYP_Thesis_Final_v2.docx'

print("Loading...")
doc = Document(PANDOC_OUT)

# ---- 1. ENHANCE TITLE PAGE ----
print("1. Enhancing title page...")
# Pandoc creates a Title paragraph. Add FYP info before it.
body = doc.element.body
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def add_para_at(doc, text, pos, size=12, bold=False, align='center'):
    """Add a paragraph at a specific position using XML."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    if bold: run.bold = True
    if align == 'center':
        p.alignment = 1
    # Move to position
    p_elem = p._element
    body.remove(p_elem)
    body.insert(pos, p_elem)
    return p

# Find the Title paragraph
title_elem = None
title_pos = 0
for i, child in enumerate(body):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        # Check if this paragraph has Title style
        pPr = child.find(f'{{{ns}}}pPr')
        if pPr is not None:
            pStyle = pPr.find(f'{{{ns}}}pStyle')
            if pStyle is not None and 'Title' in (pStyle.get(f'{{{ns}}}val') or ''):
                title_elem = child
                title_pos = i
                break

# Add FYP header paragraphs before the title
insert_pos = 0
for i, child in enumerate(body):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        insert_pos = i
        break

header_items = [
    ('CityUHK Qingdao Research Institute', 14, False),
    ('', 12, False),
    ('FINAL YEAR PROJECT REPORT', 18, True),
    ('', 12, False),
    ('2025/26-[Proposal Code]', 12, False),
    ('', 12, False),
    ('', 12, False),  # spacer
]

for text, size, bold in reversed(header_items):
    if text:
        p = doc.add_paragraph()
        p.alignment = 1
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        if bold: run.bold = True
        p_elem = p._element
        body.remove(p_elem)
        body.insert(insert_pos, p_elem)
    else:
        p = doc.add_paragraph()
        p.alignment = 1
        p_elem = p._element
        body.remove(p_elem)
        body.insert(insert_pos, p_elem)

# ---- 2. ADD STUDENT INFO AFTER TITLE ----
# Find the author paragraph after title
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == 'JU JIAXING' or 'JU JIAXING' in t:
        # Add FYP info before this
        break

# ---- 3. STYLES: ENFORCE FONTS ----
print("2. Enforcing fonts and styles...")

def set_style(name, size=None, bold=None):
    try:
        s = doc.styles[name]
    except: return
    s.font.name = 'Times New Roman'
    if size: s.font.size = Pt(size)
    if bold is not None: s.font.bold = bold
    rpr = s.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = doc.element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rf)
    for a in ['w:ascii','w:hAnsi','w:eastAsia','w:cs']:
        rf.set(qn(a), 'Times New Roman')

set_style('Normal', 12)
doc.styles['Normal'].paragraph_format.line_spacing = 1.5
set_style('Heading 1', 16, bold=True)
set_style('Heading 2', 14, bold=True)
set_style('Heading 3', 13, bold=True)
set_style('Title', 20, bold=True)

# ---- 4. ENFORCE ON ALL RUNS ----
print("3. Enforcing on all runs...")
for p in doc.paragraphs:
    try:
        runs = p.runs
    except:
        continue
    for r in runs:
        r.font.name = 'Times New Roman'
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = doc.element.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rf)
        for a in ['w:ascii','w:hAnsi','w:cs']:
            rf.set(qn(a), 'Times New Roman')

    try:
        style = p.style.name
    except:
        continue

    if style == 'Heading 1':
        for r in runs: r.font.size = Pt(16)
    elif style == 'Heading 2':
        for r in runs: r.font.size = Pt(14)
    elif style == 'Heading 3':
        for r in runs: r.font.size = Pt(13)

    if 'Heading' in style or style == 'Normal':
        p.paragraph_format.line_spacing = 1.5

    # References
    if re.match(r'^\[\d+\]', p.text.strip()[:10]):
        for r in runs: r.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5

# ---- 5. DECLARATION ----
print("4. Fixing Declaration...")
for i, p in enumerate(doc.paragraphs):
    try:
        style = p.style.name
    except:
        continue
    if style == 'Heading 1' and 'Declaration' in p.text:
        if i + 1 < len(doc.paragraphs):
            try:
                np = doc.paragraphs[i + 1]
                np.clear()
                r = np.add_run(
                    'I have read the student handbook and I understand the meaning of '
                    'academic dishonesty, in particular plagiarism and collusion. I declare '
                    'that the work submitted for the final year project does not involve '
                    'academic dishonesty. I give permission for my final year project work '
                    'to be electronically scanned and if found to involve academic dishonesty, '
                    'I am aware of the consequences as stated in the Student Handbook.'
                )
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
            except:
                pass
        break

# ---- 6. ABSTRACT ----
print("5. Verifying Abstract...")
for i, p in enumerate(doc.paragraphs):
    try:
        style = p.style.name
    except:
        continue
    if style == 'Heading 1' and 'Abstract' in p.text:
        if i + 1 < len(doc.paragraphs):
            try:
                ap = doc.paragraphs[i + 1]
                if len(ap.text) < 200:
                    print("   Abstract too short, reinserting...")
                    ap.clear()
                    r = ap.add_run(
                        'The increasing demand for energy-efficient edge AI inference has '
                        'driven interest in hardware acceleration of convolutional neural '
                        'networks (CNNs). RISC-V\'s open instruction set architecture provides '
                        'a unique opportunity to extend a standard processor with custom '
                        'instructions for domain-specific acceleration. This project presents '
                        'the design, implementation, and FPGA-based prototype validation of a '
                        'lightweight CNN accelerator integrated with a RISC-V E203 processor '
                        'through the Nuclei Instruction Co-extension (NICE) interface. The '
                        'accelerator implements a 4x4 systolic PE array supporting INT8 '
                        'quantized convolution. Six custom NICE instructions provide the '
                        'software interface. The design was validated through RTL simulation, '
                        'full-SoC simulation, and FPGA prototype bring-up on the Davinci Pro '
                        'A7-100T board. A bare-metal program was validated, ILA debugging '
                        'resolved four critical root causes, NICE test programs confirmed '
                        'correct instruction execution, and a LeNet-5 inference pipeline '
                        'achieved 70% accuracy on MNIST. This project demonstrates a complete '
                        'FPGA bring-up pipeline for a RISC-V-based CNN accelerator.'
                    )
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)
                else:
                    print(f"   Abstract OK ({len(ap.text.split())} words)")
            except:
                pass
        break

# ---- 7. CLEAN ARTIFACTS ----
print("6. Cleaning artifacts...")
for p in doc.paragraphs:
    try:
        runs = p.runs
    except:
        continue
    for r in runs:
        t = r.text
        for old, new in [('\\\\_', '_'), ('\\\\#', '#'), ('\\\\&', '&'), ('\\\\%', '%'), ('\\\\{', '{'), ('\\\\}', '}')]:
            t = t.replace(old, new)
        r.text = t

# ---- 8. MARGINS ----
print("7. Setting margins...")
for section in doc.sections:
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.5)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ---- SAVE ----
print(f"\nSaving {FINAL}...")
doc.save(FINAL)

import os
size_mb = os.path.getsize(FINAL) / (1024*1024)
print(f"Done: {size_mb:.1f}MB, {len(doc.paragraphs)} paragraphs")
