"""
Post-process Pandoc docx - v4 (NO XML manipulation, preserves images)
Only modifies styles, fonts, text content. Never touches body XML structure.
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

FINAL = r'C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex\FYP_Thesis_v3.docx'

print("Loading...")
doc = Document(r'C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex\thesis_pandoc_embed.docx')

# ---- 1. STYLES ----
print("1. Configuring styles...")

def cfg_style(name, size=None, bold=None):
    try: s = doc.styles[name]
    except: return
    s.font.name = 'Times New Roman'
    if size: s.font.size = Pt(size)
    if bold is not None: s.font.bold = bold
    rpr = s.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = doc.element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rf)
    for a in ['w:ascii','w:hAnsi','w:eastAsia','w:cs']:
        rf.set(qn(a), 'Times New Roman')

cfg_style('Normal', 12)
doc.styles['Normal'].paragraph_format.line_spacing = 1.5
cfg_style('Heading 1', 16, bold=True)
cfg_style('Heading 2', 14, bold=True)
cfg_style('Heading 3', 13, bold=True)
cfg_style('Title', 20, bold=True)

# ---- 2. ENFORCE ON EVERY RUN ----
print("2. Enforcing fonts on all runs...")
for p in doc.paragraphs:
    try: runs = p.runs
    except: continue
    try: style = p.style.name
    except: continue

    for r in runs:
        r.font.name = 'Times New Roman'
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = doc.element.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rf)
        for a in ['w:ascii','w:hAnsi','w:cs']:
            rf.set(qn(a), 'Times New Roman')

    # Size fix
    if style == 'Heading 1':
        for r in runs: r.font.size = Pt(16)
    elif style == 'Heading 2':
        for r in runs: r.font.size = Pt(14)
    elif style == 'Heading 3':
        for r in runs: r.font.size = Pt(13)

    # Line spacing
    if 'Heading' in style or style == 'Normal':
        p.paragraph_format.line_spacing = 1.5

    # References: 12pt
    if re.match(r'^\[\d+\]', p.text.strip()[:10]):
        for r in runs: r.font.size = Pt(12)

# ---- 3. ENHANCE TITLE ----
print("3. Enhancing title...")
# Find Title paragraph and make it closer to FYP format
for p in doc.paragraphs:
    try: style = p.style.name
    except: continue
    if style == 'Title':
        p.alignment = 1  # center
        for r in p.runs:
            r.font.size = Pt(18)
            r.bold = True

# ---- 4. REPLACE DECLARATION ----
print("4. Replacing Declaration...")
for i, p in enumerate(doc.paragraphs):
    try: style = p.style.name
    except: continue
    if style == 'Heading 1' and 'Declaration' in p.text:
        if i + 1 < len(doc.paragraphs):
            try:
                np = doc.paragraphs[i + 1]
                np.clear()
                r = np.add_run(
                    'I have read the student handbook and I understand the meaning of '
                    'academic dishonesty, in particular plagiarism and collusion. I declare '
                    'that the work submitted for the final year project does not involve '
                    'academic dishonesty. I give permission for my final year project work '
                    'to be electronically scanned and if found to involve academic dishonesty, '
                    'I am aware of the consequences as stated in the Student Handbook.')
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
            except: pass
        break

# ---- 5. ENSURE ABSTRACT ----
print("5. Ensuring Abstract...")
for i, p in enumerate(doc.paragraphs):
    try: style = p.style.name
    except: continue
    if style == 'Heading 1' and 'Abstract' in p.text:
        if i + 1 < len(doc.paragraphs):
            try:
                ap = doc.paragraphs[i + 1]
                if len(ap.text.strip()) < 200:
                    ap.clear()
                    r = ap.add_run(
                        'The increasing demand for energy-efficient edge AI inference has '
                        'driven interest in hardware acceleration of convolutional neural '
                        'networks (CNNs). RISC-V\'s open instruction set architecture provides '
                        'a unique opportunity to extend a standard processor with custom '
                        'instructions for domain-specific acceleration. This project presents '
                        'the design, implementation, and FPGA-based prototype validation of a '
                        'lightweight CNN accelerator integrated with a RISC-V E203 processor '
                        'through the Nuclei Instruction Co-extension (NICE) interface. The '
                        'accelerator implements a 4x4 systolic PE array supporting INT8 '
                        'quantized convolution, with six custom NICE instructions. The design '
                        'was validated through RTL simulation, full-SoC simulation, and FPGA '
                        'prototype bring-up on the Davinci Pro A7-100T board. A bare-metal '
                        'program was validated on the FPGA, ILA debugging resolved four '
                        'critical root causes, NICE test programs confirmed correct instruction '
                        'execution, and an end-to-end LeNet-5 inference pipeline achieved 70% '
                        'accuracy on MNIST. This project demonstrates a complete FPGA bring-up '
                        'pipeline for a RISC-V-based CNN accelerator.')
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)
            except: pass
        break

# ---- 6. CLEAN ARTIFACTS ----
print("6. Cleaning LaTeX artifacts...")
for p in doc.paragraphs:
    try: runs = p.runs
    except: continue
    for r in runs:
        t = r.text
        for old, new in [('\\\\_', '_'), ('\\\\#', '#'), ('\\\\&', '&'),
                         ('\\\\%', '%'), ('\\\\{', '{'), ('\\\\}', '}')]:
            t = t.replace(old, new)
        r.text = t

# ---- 7. MARGINS ----
print("7. Setting margins...")
for section in doc.sections:
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.5)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ---- SAVE ----
print(f"Saving {FINAL}...")
doc.save(FINAL)

# Verify
doc2 = Document(FINAL)
drawings = len(doc2.element.findall(
    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
size_mb = os.path.getsize(FINAL) / (1024*1024)
print(f"Done: {size_mb:.1f}MB, {len(doc2.paragraphs)} paras, {drawings} images")
