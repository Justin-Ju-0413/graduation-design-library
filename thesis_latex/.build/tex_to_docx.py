"""
Convert LaTeX thesis to FYP Word template.
Reads chapters from thesis_latex/ and writes into the docx template.
"""
import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

LATEX_DIR = r"C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex"
FIG_DIR = os.path.join(LATEX_DIR, "figures")
TEMPLATE = r"C:\Users\16084\Documents\Graduation_Design_Library\11_FYP_requirement\3+1+X FYP_Report_Template_251024.docx"
OUTPUT = os.path.join(LATEX_DIR, "FYP_Thesis_Final.docx")

# Read LaTeX chapter
def read_tex(filename):
    path = os.path.join(LATEX_DIR, "chapters", filename)
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

# Strip LaTeX commands for plain text (keep structure)
def clean_tex(text):
    # Remove comments
    text = re.sub(r'(?<!\\)%.*$', '', text, flags=re.MULTILINE)
    # Remove figure/table environments (we insert figures separately)
    text = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', '', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{table\}.*?\\end\{table\}', '', text, flags=re.DOTALL)
    # Replace \section{...} → heading markers
    text = re.sub(r'\\section\{([^}]*)\}', r'[H2]\1[/H2]', text)
    text = re.sub(r'\\subsection\{([^}]*)\}', r'[H3]\1[/H3]', text)
    text = re.sub(r'\\subsubsection\{([^}]*)\}', r'[H4]\1[/H4]', text)
    # Replace \cite{...} → [refs]
    text = re.sub(r'~?\\cite\{([^}]*)\}', r' [\1] ', text)
    # Replace \ref{...} → X
    text = re.sub(r'~?\\ref\{([^}]*)\}', r' \1 ', text)
    # Replace \label{...}
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    # Replace inline commands
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\code\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\reg\{([^}]*)\}', r'\1', text)
    # Replace \emph
    text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
    # Replace math mode
    text = re.sub(r'\$([^$]*)\$', r'\1', text)
    text = re.sub(r'\\\[.*?\\\]', '', text, flags=re.DOTALL)
    # Replace \textasciitilde, \textasciicircum
    text = text.replace('\\textasciitilde{}', '~')
    text = text.replace('\\textasciicircum{}', '^')
    # Replace times symbol
    text = text.replace('$\\times$', 'x')
    text = text.replace('\\times', 'x')
    # Replace \_ and \#
    text = text.replace('\\_', '_')
    text = text.replace('\\#', '#')
    # Replace itemize
    text = text.replace('\\begin{itemize}', '')
    text = text.replace('\\end{itemize}', '')
    text = re.sub(r'\\item\s*', '  • ', text)
    # Remove lstlisting blocks
    text = re.sub(r'\\begin\{lstlisting\}.*?\\end\{lstlisting\}', '', text, flags=re.DOTALL)
    # Clean up multiple newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Clean up trailing spaces
    text = re.sub(r' +', ' ', text)
    return text.strip()

# Split cleaned text into paragraphs
def text_to_paragraphs(text):
    paras = []
    for line in text.split('\n\n'):
        line = line.strip()
        if not line:
            continue
        # Check for heading markers
        if line.startswith('[H2]') and line.endswith('[/H2]'):
            h = line[4:-5].strip()
            paras.append(('h2', h))
        elif line.startswith('[H3]') and line.endswith('[/H3]'):
            h = line[4:-5].strip()
            paras.append(('h3', h))
        elif line.startswith('[H4]') and line.endswith('[/H4]'):
            h = line[4:-5].strip()
            paras.append(('h4', h))
        else:
            # Remove inline heading markers
            clean = line.replace('[H2]', '').replace('[/H2]', '')
            clean = clean.replace('[H3]', '').replace('[/H3]', '')
            clean = clean.replace('[H4]', '').replace('[/H4]', '')
            clean = re.sub(r'^\s*•\s*', '• ', clean)
            if clean.strip():
                paras.append(('p', clean.strip()))
    return paras

# Add paragraph to document
def add_para(doc, text, style='Normal', bold=False, size=12, alignment=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if bold:
        run.bold = True
    if alignment is not None:
        p.alignment = alignment
    return p

def add_heading_styled(doc, text, level=1):
    # Template only has 'Heading 1' and 'Heading 2'. Map h3/h4 → bold paragraph
    if level <= 2:
        try:
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.name = 'Times New Roman'
            return h
        except KeyError:
            pass
    # Fallback: bold paragraph as sub-heading
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13) if level == 2 else Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    return p


# ================================================================
# MAIN
# ================================================================
print("Opening template...")
doc = Document(TEMPLATE)

# Clear all existing content except the first few cover page elements
# Strategy: find where sections start and clear from there
body = doc.element.body
paras_to_remove = []
removing = False
for i, p in enumerate(doc.paragraphs):
    if '1. Introduction' in p.text and 'Checklist' not in p.text and 'e.g.' not in p.text:
        # This is the Introduction heading — keep it but remove all after
        # Actually, let's just clear all content after Abstract
        pass

# Simpler approach: work with the template by finding and replacing key sections
# Clear all paragraphs after the TOC
print("Processing template structure...")

# Find paragraph indices for key sections
intro_idx = None
methodology_idx = None
results_idx = None
discussion_idx = None
conclusion_idx = None
appendices_idx = None
references_idx = None
checklist_start = None

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == '1. Introduction':
        intro_idx = i
    elif t.startswith('2. Methodology'):
        methodology_idx = i
    elif t.startswith('3. Results'):
        results_idx = i
    elif t.startswith('4. Discussion'):
        discussion_idx = i
    elif t.startswith('5. Conclusion'):
        conclusion_idx = i
    elif t.startswith('Appendices'):
        appendices_idx = i
    elif t.startswith('References:'):
        references_idx = i

print(f"Found sections: Intro={intro_idx}, Meth={methodology_idx}, Results={results_idx}, "
      f"Disc={discussion_idx}, Concl={conclusion_idx}, App={appendices_idx}, Ref={references_idx}")

# Fill cover page fields
print("\nFilling cover page...")
# Paragraph 17: Student Name
doc.paragraphs[17].clear()
doc.paragraphs[17].add_run('Student Name: JU JIAXING').font.size = Pt(12)
# Paragraph 18: Home University
doc.paragraphs[18].clear()
doc.paragraphs[18].add_run('Home University: City University of Hong Kong').font.size = Pt(12)
# Paragraph 19: Student ID
doc.paragraphs[19].clear()
doc.paragraphs[19].add_run('Student ID: [Please fill in]').font.size = Pt(12)
# Paragraph 20: Supervisor
doc.paragraphs[20].clear()
doc.paragraphs[20].add_run('Supervisor: [Supervisor Name]').font.size = Pt(12)
# Paragraph 21: Assessor
doc.paragraphs[21].clear()
doc.paragraphs[21].add_run('Assessor: [Assessor Name]').font.size = Pt(12)
# Paragraph 11: Proposal Code
doc.paragraphs[11].clear()
doc.paragraphs[11].add_run('2025/26-[Proposal Code]').font.size = Pt(12)
# Paragraph 13: Project Title
doc.paragraphs[13].clear()
title_run = doc.paragraphs[13].add_run('RISC-V Custom Instruction Based Lightweight CNN Accelerator FPGA Prototype Validation')
title_run.font.size = Pt(14)
title_run.bold = True

# Fill Declaration (paragraph 28)
print("Filling Declaration...")
doc.paragraphs[28].clear()
doc.paragraphs[28].add_run(
    'I have read the student handbook and I understand the meaning of academic dishonesty, '
    'in particular plagiarism and collusion. I declare that the work submitted for the final '
    'year project does not involve academic dishonesty. I give permission for my final year '
    'project work to be electronically scanned and if found to involve academic dishonesty, '
    'I am aware of the consequences as stated in the Student Handbook.'
).font.size = Pt(12)

# Fill Abstract (paragraph 44-54 area)
print("Filling Abstract...")
abstract_text = read_tex('00_abstract.tex')
abstract_clean = clean_tex(abstract_text)
# Clear checklist paragraphs
for i in range(44, 55):
    if i < len(doc.paragraphs):
        doc.paragraphs[i].clear()

doc.paragraphs[44].add_run('Abstract').bold = True
doc.paragraphs[45].clear()
doc.paragraphs[45].add_run(abstract_clean).font.size = Pt(12)

# Remove checklist content in intro/methodology/etc sections
# Fill content sections
print("Filling Introduction...")
# Clear checklist items between Introduction heading and 1.1 Background
for i in range(intro_idx+1, intro_idx+12):
    if i < len(doc.paragraphs):
        doc.paragraphs[i].clear()

# Read all chapter content
ch1 = read_tex('01_1_introduction.tex')
ch2 = read_tex('02_2_background.tex')
ch3 = read_tex('03_3_methodology.tex')
ch4 = read_tex('04_4_results.tex')
ch5 = read_tex('05_5_discussion.tex')
ch6 = read_tex('06_6_conclusion.tex')

# Clean and parse
ch1_paras = text_to_paragraphs(clean_tex(ch1))
ch2_paras = text_to_paragraphs(clean_tex(ch2))
ch3_paras = text_to_paragraphs(clean_tex(ch3))
ch4_paras = text_to_paragraphs(clean_tex(ch4))
ch5_paras = text_to_paragraphs(clean_tex(ch5))
ch6_paras = text_to_paragraphs(clean_tex(ch6))

# Helper to insert content after a specific paragraph
def insert_content_after(doc, after_idx, paras_list):
    """Insert content paragraphs after the given paragraph index."""
    # We insert at the end since python-docx can't easily insert at a specific position
    pass

# Since python-docx adds at end, let's take a different approach:
# Clear all template guide text from Introduction onwards
# Then add all content at the end

print("Clearing template guide text...")
# Remove paragraphs from Introduction heading to end (keep cover, declaration, abstract)
clear_from = intro_idx
paras_to_keep = []
for i, p in enumerate(doc.paragraphs):
    if i >= clear_from:
        p.clear()

# Now add content sections as new paragraphs after the last kept one
# First add a page break before main content
print("Adding content...")

# ---- 1. Introduction ----
add_heading_styled(doc, '1. Introduction', level=1)
for ptype, text in ch1_paras:
    if ptype == 'h2':
        add_heading_styled(doc, text, level=2)
    elif ptype == 'h3':
        add_heading_styled(doc, text, level=3)
    else:
        add_para(doc, text, size=12)

# ---- 2. Methodology ----
add_heading_styled(doc, '2. Methodology', level=1)
for ptype, text in ch3_paras:
    if ptype == 'h2':
        add_heading_styled(doc, text, level=2)
    elif ptype == 'h3':
        add_heading_styled(doc, text, level=3)
    else:
        add_para(doc, text, size=12)

# ---- 3. Results ----
add_heading_styled(doc, '3. Results', level=1)
for ptype, text in ch4_paras:
    if ptype == 'h2':
        add_heading_styled(doc, text, level=2)
    elif ptype == 'h3':
        add_heading_styled(doc, text, level=3)
    else:
        add_para(doc, text, size=12)

# ---- 4. Discussion ----
add_heading_styled(doc, '4. Discussion', level=1)
for ptype, text in ch5_paras:
    if ptype == 'h2':
        add_heading_styled(doc, text, level=2)
    elif ptype == 'h3':
        add_heading_styled(doc, text, level=3)
    else:
        add_para(doc, text, size=12)

# ---- 5. Conclusion ----
add_heading_styled(doc, '5. Conclusion', level=1)
for ptype, text in ch6_paras:
    if ptype == 'h2':
        add_heading_styled(doc, text, level=2)
    elif ptype == 'h3':
        add_heading_styled(doc, text, level=3)
    else:
        add_para(doc, text, size=12)

# ---- References ----
add_heading_styled(doc, 'References', level=1)

# Add references from .bib file
bib_path = os.path.join(LATEX_DIR, 'references.bib')
with open(bib_path, 'r', encoding='utf-8') as f:
    bib_content = f.read()

# Parse bib entries
bib_entries = []
for match in re.finditer(r'@(\w+)\{([^,]+),\s*\n(.*?)\n\}', bib_content, re.DOTALL):
    etype = match.group(1)
    key = match.group(2)
    fields = match.group(3)

    title = re.search(r'title\s*=\s*\{([^}]*)\}', fields)
    author = re.search(r'author\s*=\s*\{([^}]*)\}', fields)
    journal = re.search(r'journal\s*=\s*\{([^}]*)\}', fields)
    booktitle = re.search(r'booktitle\s*=\s*\{([^}]*)\}', fields)
    year = re.search(r'year\s*=\s*\{([^}]*)\}', fields)
    doi = re.search(r'doi\s*=\s*\{([^}]*)\}', fields)
    volume = re.search(r'volume\s*=\s*\{([^}]*)\}', fields)
    number = re.search(r'number\s*=\s*\{([^}]*)\}', fields)
    pages = re.search(r'pages\s*=\s*\{([^}]*)\}', fields)
    url = re.search(r'url\s*=\s*\{([^}]*)\}', fields)
    institution = re.search(r'institution\s*=\s*\{([^}]*)\}', fields)
    school = re.search(r'school\s*=\s*\{([^}]*)\}', fields)
    publisher = re.search(r'publisher\s*=\s*\{([^}]*)\}', fields)
    note = re.search(r'note\s*=\s*\{([^}]*)\}', fields)

    t = title.group(1) if title else ''
    a = author.group(1) if author else ''
    j = journal.group(1) if journal else ''
    bt = booktitle.group(1) if booktitle else ''
    y = year.group(1) if year else ''
    v = volume.group(1) if volume else ''
    n = number.group(1) if number else ''
    pp = pages.group(1) if pages else ''
    inst = institution.group(1) if institution else ''
    sch = school.group(1) if school else ''
    pub = publisher.group(1) if publisher else ''
    nt = note.group(1) if note else ''

    bib_entries.append({
        'key': key, 'type': etype, 'title': t, 'author': a,
        'journal': j, 'booktitle': bt, 'year': y,
        'volume': v, 'number': n, 'pages': pp,
        'institution': inst, 'school': sch, 'publisher': pub, 'note': nt
    })

ref_num = 1
for entry in bib_entries:
    # Format in IEEE style
    ref_text = f"[{ref_num}] "
    if entry['author']:
        ref_text += entry['author'] + ', '
    ref_text += f'"{entry["title"]}," '
    if entry['journal']:
        ref_text += f'{entry["journal"]}, '
    elif entry['booktitle']:
        ref_text += f'in {entry["booktitle"]}, '
    elif entry['institution']:
        ref_text += f'{entry["institution"]}, '
    elif entry['school']:
        ref_text += f'{entry["school"]}, '
    elif entry['publisher']:
        ref_text += f'{entry["publisher"]}, '

    if entry['volume']:
        ref_text += f'vol. {entry["volume"]}, '
    if entry['number']:
        ref_text += f'no. {entry["number"]}, '
    if entry['pages']:
        ref_text += f'pp. {entry["pages"]}, '
    if entry['year']:
        ref_text += f'{entry["year"]}.'
    if entry['note']:
        ref_text += f' {entry["note"]}.'

    add_para(doc, ref_text, size=10)
    ref_num += 1

# Save
print(f"\nSaving to {OUTPUT}...")
doc.save(OUTPUT)
print(f"Done! File: {OUTPUT}")
print(f"Content: {len(ch1_paras)} intro + {len(ch3_paras)} methodology + "
      f"{len(ch4_paras)} results + {len(ch5_paras)} discussion + "
      f"{len(ch6_paras)} conclusion + {len(bib_entries)} references")
