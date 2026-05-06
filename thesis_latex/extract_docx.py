#!/usr/bin/env python3
"""Extract thesis content from .docx for LaTeX migration."""

import sys
import os
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import json

DOCX_PATH = r"C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing\FYP_Thesis_Final_v2.docx"
OUTPUT_DIR = r"C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex"

def get_style_name(paragraph):
    """Get the style name of a paragraph."""
    if paragraph.style:
        return paragraph.style.name or ""
    return ""

def get_run_formatting(paragraph):
    """Extract formatting info from runs."""
    info = []
    for run in paragraph.runs:
        bold = run.bold
        italic = run.italic
        font_name = run.font.name
        font_size = run.font.size
        text = run.text
        info.append({
            'text': text,
            'bold': bold,
            'italic': italic,
            'font_name': font_name,
            'font_size': str(font_size) if font_size else None,
        })
    return info

def extract_tables(doc):
    """Extract all tables from the document."""
    tables = []
    for i, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        # Determine if header row exists
        has_header = False
        if len(rows_data) > 1:
            # Check if first row is bold/header-like
            first_row_cells = table.rows[0].cells
            for cell in first_row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.bold:
                            has_header = True
                            break

        tables.append({
            'index': i,
            'rows': rows_data,
            'nrows': len(rows_data),
            'ncols': max(len(r) for r in rows_data) if rows_data else 0,
            'has_header': has_header,
        })
    return tables

def extract_document():
    """Extract full document structure."""
    doc = Document(DOCX_PATH)

    output = {
        'paragraphs': [],
        'tables': [],
        'sections': [],
    }

    current_section = None

    for i, para in enumerate(doc.paragraphs):
        style = get_style_name(para)
        text = para.text.strip()
        runs_info = get_run_formatting(para)

        p_info = {
            'index': i,
            'style': style,
            'text': text,
            'runs': runs_info,
            'alignment': str(para.alignment) if para.alignment else None,
        }

        output['paragraphs'].append(p_info)

        # Track section structure
        if style.startswith('Heading 1') or style.startswith('heading 1'):
            current_section = {
                'title': text,
                'start_para': i,
                'subsections': []
            }
            output['sections'].append(current_section)
        elif style.startswith('Heading 2') or style.startswith('heading 2'):
            if current_section is not None:
                current_section['subsections'].append({
                    'title': text,
                    'start_para': i,
                })

    output['tables'] = extract_tables(doc)

    # Save intermediate data
    json_path = os.path.join(OUTPUT_DIR, 'docx_structure.json')
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"Extracted {len(output['paragraphs'])} paragraphs, {len(output['tables'])} tables")
    print(f"Found {len(output['sections'])} sections")
    for s in output['sections']:
        print(f"  Section: {s['title']} ({len(s['subsections'])} subsections)")

    return output

def print_paragraphs_by_style(doc_data):
    """Print first few paragraphs of each style."""
    styles_seen = {}
    for p in doc_data['paragraphs'][:200]:
        s = p['style']
        if s not in styles_seen:
            styles_seen[s] = p['text'][:120]

    print("\n=== Paragraph Styles Found ===")
    for s, t in styles_seen.items():
        print(f"  '{s}': '{t}'")

if __name__ == '__main__':
    data = extract_document()
    print_paragraphs_by_style(data)

    # Print table info
    print(f"\n=== Tables: {len(data['tables'])} ===")
    for t in data['tables']:
        print(f"  Table {t['index']}: {t['nrows']}x{t['ncols']} (header: {t['has_header']})")
        for r in t['rows'][:3]:
            print(f"    Row: {[c[:40] for c in r]}")
        if len(t['rows']) > 3:
            print(f"    ... ({len(t['rows'])-3} more rows)")
