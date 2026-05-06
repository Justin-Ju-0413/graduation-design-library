#!/usr/bin/env python3
"""
Comprehensive .docx to LaTeX thesis migration script.
Processes FYP_Thesis_Final_v2.docx and generates complete LaTeX files.
"""

import json, os, re, shutil
from docx import Document

DOCX_PATH = r"C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing\FYP_Thesis_Final_v2.docx"
OUTPUT_DIR = r"C:\Users\16084\Documents\Graduation_Design_Library\thesis_latex"
CHAPTERS_DIR = os.path.join(OUTPUT_DIR, "chapters")
FIGURES_SRC = r"C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing\Figures"
FIGURES_DST = os.path.join(OUTPUT_DIR, "figures")

os.makedirs(CHAPTERS_DIR, exist_ok=True)
os.makedirs(FIGURES_DST, exist_ok=True)

# Load pre-extracted JSON
with open(os.path.join(OUTPUT_DIR, 'docx_structure.json'), 'r', encoding='utf-8') as f:
    doc_data = json.load(f)
paragraphs = doc_data['paragraphs']

doc = Document(DOCX_PATH)

# Chapter boundaries (Heading 1 indices from Final_v2)
HEADING1_CHAPTERS = {
    'Declaration': 1, 'Abstract': 8,
    'Chapter 1. Introduction': 15, 'Chapter 2. Background': 41,
    'Chapter 3. Methodology': 192, 'Chapter 4. Results': 332,
    'Chapter 5. Discussion': 404, 'Chapter 6. Conclusion': 437,
    'References': 460,
}

# Reference number -> BibTeX key
REF_MAP = {
    '1': 'RISCV_UNPRIV_SPEC', '2': 'WATERMAN_PHD', '3': 'ASANOVIC_HOTCHIPS',
    '4': 'RISCV_READER', '5': 'RISCV_PRIV_SPEC', '6': 'NUCLEI_E203',
    '7': 'NUCLEI_SDK', '8': 'KUNG_SYSTOLIC', '9': 'SZE_SURVEY',
    '10': 'LENET', '11': 'ALEXNET', '12': 'JACOB_QUANT',
    '13': 'KRISHNAMOORTHI_QUANT', '14': 'HAN_DEEPCOMPRESS', '15': 'EYERISS',
    '16': 'TPU', '17': 'DIANNAO', '18': 'GEMMINI',
    '19': 'XILINX_UG908', '20': 'XILINX_UG901', '21': 'XILINX_UG904',
    '22': 'KUON_FPGA_GAP', '23': 'NUCLEI_NICE', '24': 'ROCKETCHIP',
    '25': 'GDB_MANUAL', '26': 'OPENOCD_MANUAL', '27': 'HAN_DEEPCOMPRESS',
}

FIG_FILES = {
    '3.1': 'fig3_1_soc_architecture.png', '3.2': 'fig3_2_instruction_format.png',
    '3.2b': 'fig3_2b_instruction_table.png', '3.3': 'fig3_3_pe_microarchitecture.png',
    '3.4': 'fig3_4_pe_array.png', '3.5': 'fig3_5_packed_format.png',
    '3.6': 'fig3_6_build_pipeline.png', '4.1': 'fig_ila_pc_trace.png',
    '4.2': 'fig_ila_nice_activity.png', '4.3': 'fig_speedup_bar.png',
    '4.4': 'fig_resource_pie.png', '4.5': 'fig_timing.png',
    '4.6': 'fig_utilization.png',
}

# Helpers
LATEX_SPECIAL = {'\\': '\\textbackslash{}', '{': '\\{', '}': '\\}',
                  '$': '\\$', '&': '\\&', '#': '\\#', '_': '\\_',
                  '%': '\\%', '~': '\\textasciitilde{}', '^': '\\textasciicircum{}'}

def escape_latex(text):
    parts = re.split(r'(\$[^$]*\$)', text)
    result = []
    for part in parts:
        if part.startswith('$') and part.endswith('$'):
            result.append(part)
        else:
            for old, new in LATEX_SPECIAL.items():
                part = part.replace(old, new)
            result.append(part)
    return ''.join(result)

def format_citation(text):
    def replace_ref(m):
        nums = re.findall(r'\d+', m.group(1))
        keys = [REF_MAP[n] for n in nums if n in REF_MAP]
        return '\\cite{' + ','.join(keys) + '}' if keys else m.group(0)
    return re.sub(r'\[(\d+(?:\s*[,]\s*\d+)*)\]', replace_ref, text)

def is_table_caption(text):
    return bool(re.match(r'Table\s+\d+\.\d+\s*:', text))

def is_figure_caption(text):
    return bool(re.match(r'Figure\s+\d+\.\d+\s*:', text))

def is_ascii_diagram(line):
    if not line.strip():
        return False
    return sum(1 for c in line if c in '+-|*') >= 3 and len(line.strip()) > 15

def detect_lang(code_lines):
    first = '\n'.join(code_lines[:5])
    if 'module ' in first or 'endmodule' in first:
        return '[language=Verilog,style=verilog]'
    if '#define ' in first or '__asm__' in first or '#include' in first:
        return '[language=C,style=cstyle]'
    return '[style=ascii]'

# Table mapping
def build_table_mapping():
    mapping = {}
    docx_ti = 0
    claimed = set()

    # Phase 1: match explicit "Table X.Y:" captions
    for i, p in enumerate(paragraphs):
        text = p['text']
        if not is_table_caption(text):
            continue
        m = re.match(r'Table\s+(\d+\.\d+):\s*(.*)', text)
        if not m:
            continue
        tbl_num = m.group(1)
        if tbl_num == '3.4':  # ASCII art
            mapping[i] = (tbl_num, None, m.group(2))
            continue
        if docx_ti < len(doc.tables):
            mapping[i] = (tbl_num, docx_ti, m.group(2))
            claimed.add(docx_ti)
            docx_ti += 1

    # Phase 2: unclaimed tables
    unclaimed = [
        (7, 'The simulation required several configuration fixes',
         'Summary of build configuration fixes for FPGA compatibility'),
        (8, 'The FPGA build flow uses Xilinx Vivado',
         'FPGA build modes and their descriptions'),
        (9, 'All builds achieved clean timing closure',
         'Timing closure results across FPGA build configurations'),
        (10, 'The ILA capture.*showed the CPU executing',
         'hello_e203 ILA probe values'),
        (11, 'The ILA capture confirmed that the CPU executed through all NICE',
         'NICE accelerator ILA probe values'),
        (12, 'FPGA resource utilization.*summarized below',
         'FPGA resource utilization for the complete SoC'),
        (13, 'Comparison with Project Objectives',
         'Project objective status summary'),
    ]
    ch4_cnt, ch5_cnt = 1, 1
    for dx, pattern, caption in unclaimed:
        if dx in claimed:
            continue
        for i, p in enumerate(paragraphs):
            if re.search(pattern, p['text'], re.IGNORECASE):
                ch_num = 4
                for j in range(i, -1, -1):
                    h1 = paragraphs[j]
                    if h1['style'] == 'Heading 1':
                        m = re.match(r'Chapter\s+(\d+)', h1['text'])
                        if m:
                            ch_num = int(m.group(1))
                        break
                tbl_num = f'{ch_num}.{ch4_cnt}' if ch_num == 4 else f'{ch_num}.{ch5_cnt}'
                if ch_num == 4:
                    ch4_cnt += 1
                else:
                    ch5_cnt += 1
                mapping[i] = (tbl_num, dx, caption)
                claimed.add(dx)
                print(f"  Table {tbl_num} -> docx[{dx}] at para[{i}]")
                break
    return mapping

TABLE_MAP = build_table_mapping()
print(f"Total mapped tables: {len(TABLE_MAP)}")

def extract_table_latex(idx):
    if idx is None or idx >= len(doc.tables):
        return None
    table = doc.tables[idx]
    rows = []
    for row in table.rows:
        rows.append([c.text.strip() for c in row.cells])
    if not rows:
        return None
    ncols = max(len(r) for r in rows)
    fmt = 'l' + 'c' * (ncols - 1)
    latex = []
    for ri, row in enumerate(rows):
        while len(row) < ncols:
            row.append('')
        line = ' & '.join(escape_latex(c) for c in row) + ' \\\\'
        line += ' \\hline\\hline' if ri == 0 else ' \\hline'
        latex.append(line)
    return '\\begin{tabular}{' + fmt + '}\n\\hline\n' + '\n'.join(latex) + '\n\\end{tabular}'

def emit_table(docx_idx, tbl_num, caption_text, lines):
    tabular = extract_table_latex(docx_idx)
    if not tabular:
        return
    label = f'tab:{tbl_num.replace(".", "_")}'
    lines.append('\\begin{table}[htbp]')
    lines.append('\\centering')
    lines.append('\\caption{' + escape_latex(caption_text) + '}')
    lines.append('\\label{' + label + '}')
    lines.append(tabular)
    lines.append('\\end{table}')
    lines.append('')

# ======================================================================
# Chapter generation
# ======================================================================

def get_chapter_range(title):
    titles = list(HEADING1_CHAPTERS.keys())
    if title not in HEADING1_CHAPTERS:
        return None, None
    start = HEADING1_CHAPTERS[title]
    idx = titles.index(title)
    end = HEADING1_CHAPTERS[titles[idx + 1]] if idx + 1 < len(titles) else len(paragraphs)
    return start, end

def generate_chapter(ch_num, ch_title):
    start, end = get_chapter_range(ch_title)
    if start is None:
        return

    name = ch_title.lower().replace('chapter ', '').replace('. ', '_')
    fpath = os.path.join(CHAPTERS_DIR, f'{ch_num:02d}_{name}.tex')

    lines = [f'% Chapter {ch_num}: {ch_title}', '', '']
    src_buf = []
    in_src = False

    def flush_src():
        nonlocal src_buf, in_src
        if not src_buf:
            return
        if sum(1 for l in src_buf if is_ascii_diagram(l)) >= 2:
            lines.append('\\begin{lstlisting}[style=ascii]')
            for l in src_buf:
                lines.append(l.rstrip())
            lines.append('\\end{lstlisting}')
        else:
            lines.append('\\begin{lstlisting}' + detect_lang(src_buf))
            for l in src_buf:
                lines.append(l.rstrip())
            lines.append('\\end{lstlisting}')
        lines.append('')
        src_buf = []
        in_src = False

    # Determine which tables fall in this chapter and their target paragraph indices
    chap_tables = {idx: info for idx, info in TABLE_MAP.items() if start <= idx < end}

    for i in range(start, end):
        if i >= len(paragraphs):
            break
        p = paragraphs[i]
        text = p['text']
        style = p['style']

        # Source code blocks
        if style == 'Source Code':
            if not in_src:
                flush_src()
                in_src = True
            src_buf.append(text)
            continue
        if in_src:
            flush_src()

        # Skip empty paragraphs (but still check for tables at their index)
        if not text:
            if i in chap_tables:
                tbl_num, dx, caption = chap_tables[i]
                if dx is not None:
                    emit_table(dx, tbl_num, caption, lines)
            continue

        # Heading 1
        if style == 'Heading 1':
            continue

        # ====== Process paragraph content ======

        # Heading 2 -> section
        if style == 'Heading 2':
            flush_src()
            m = re.match(r'(\d+\.\d+)\s+(.*)', text)
            if m:
                lines.append(f'\\section{{{escape_latex(m.group(2))}}}\\label{{sec:{m.group(1).replace(".", "_")}}}')
            else:
                lines.append(f'\\section{{{escape_latex(text)}}}')
            lines.append('')
            if i in chap_tables:
                tbl_num, dx, caption = chap_tables[i]
                if dx is not None:
                    emit_table(dx, tbl_num, caption, lines)
            continue

        # Figure caption
        if is_figure_caption(text):
            flush_src()
            m = re.match(r'Figure\s+(\d+\.\d+):?\s*(.*)', text)
            if m:
                fig_num, fig_desc = m.group(1), m.group(2)
                lines.append('\\begin{figure}[htbp]')
                lines.append('\\centering')
                img = FIG_FILES.get(fig_num, '')
                if img:
                    lines.append('\\includegraphics[width=0.85\\textwidth]{figures/' + img + '}')
                lines.append('\\caption{' + escape_latex(fig_desc) + '}')
                lines.append('\\label{' + f'fig:{fig_num.replace(".", "_")}' + '}')
                lines.append('\\end{figure}')
                lines.append('')
            continue

        # Regular text
        if style in ('Normal', 'Body Text'):
            # Subsubsection headings (2.1.1, 3.2.1, etc.)
            sub_m = re.match(r'(\d+\.\d+\.\d+)\s+(.*)', text)
            if sub_m:
                lines.append(f'\\subsection{{{escape_latex(sub_m.group(2))}}}')
                lines.append('')
            elif re.match(r'Stage\s+\d+:', text):
                lines.append(f'\\subsubsection{{{escape_latex(text)}}}')
                lines.append('')
            elif any(text.strip().startswith(h) and text.strip().endswith(':')
                     for h in ['Interface Signals', 'Handshake Protocol',
                        'Integration in the Subsystem', 'Instruction Encoding Convention',
                        'Request Channel', 'Response Channel', 'Memory Channel',
                        'Request channel', 'Response channel',
                        'Memory request channel', 'Memory response channel']):
                lines.append(f'\\subsection*{{{escape_latex(text)}}}')
                lines.append('')
            else:
                clean = format_citation(text)
                clean = re.sub(r'\s*# Chapter \d+\.?\s*.*', '', clean)
                lines.append(escape_latex(clean))
                lines.append('')
                lines.append('')

            # Check for table at this paragraph index (emit AFTER text)
            if i in chap_tables:
                tbl_num, dx, caption = chap_tables[i]
                if dx is not None:
                    emit_table(dx, tbl_num, caption, lines)
            continue

    if in_src:
        flush_src()

    with open(fpath, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print(f"  {os.path.basename(fpath)} ({len(lines)} lines)")

def generate_declaration():
    start, end = get_chapter_range('Declaration')
    lines = ['% Declaration']
    for i in range(start + 1, end):
        t = paragraphs[i]['text']
        if t:
            lines.append(escape_latex(t))
            lines.append('')
    if len(lines) < 3:
        lines = [
            'I hereby declare that this thesis is my own original work.',
            '',
            '\\vspace{1cm}',
            '\\noindent Student Name: \\rule{6cm}{0.4pt} \\hfill Date: \\rule{4cm}{0.4pt}',
            '\\noindent Signature: \\rule{6cm}{0.4pt}',
        ]
    with open(os.path.join(CHAPTERS_DIR, '00_declaration.tex'), 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print("  00_declaration.tex")

def generate_abstract():
    start, end = get_chapter_range('Abstract')
    lines = ['% Abstract']
    for i in range(start + 1, end):
        t = paragraphs[i]['text']
        if t:
            lines.append(escape_latex(format_citation(t)))
            lines.append('')
    with open(os.path.join(CHAPTERS_DIR, '00_abstract.tex'), 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print("  00_abstract.tex")

def generate_main():
    content = r"""%===============================================================================
% RISC-V Custom Instruction Based Lightweight CNN Accelerator
% FPGA Prototype Validation - Main LaTeX File
% XeLaTeX + Biber
%===============================================================================

\documentclass[12pt,a4paper,twoside]{report}

\usepackage[top=1.5in, bottom=1.5in, left=1in, right=1in]{geometry}
\usepackage{setspace}
\onehalfspacing

\usepackage{fontspec}
\setmainfont{Times New Roman}
\usepackage{xeCJK}
\setCJKmainfont{SimSun}

\usepackage{amsmath,amssymb,mathtools}
\usepackage{graphicx}
\graphicspath{{figures/}}

\usepackage{booktabs,tabularx,array,longtable,multirow}
\usepackage[font=small,labelfont=bf]{caption}
\usepackage{subcaption}

\usepackage{fancyhdr}
\pagestyle{fancy}
\fancyhf{}
\fancyhead[LE,RO]{\thepage}
\fancyhead[RE]{\leftmark}
\fancyhead[LO]{\rightmark}
\renewcommand{\headrulewidth}{0.4pt}
\renewcommand{\chaptermark}[1]{\markboth{\MakeUppercase{#1}}{}}
\renewcommand{\sectionmark}[1]{\markright{\thesection\ #1}}

\usepackage[colorlinks=true,linkcolor=black,citecolor=black,urlcolor=blue]{hyperref}
\usepackage[titles]{tocloft}
\setcounter{tocdepth}{2}
\setcounter{secnumdepth}{3}

\usepackage{listings,xcolor}
\definecolor{codegreen}{rgb}{0,0.6,0}
\definecolor{codegray}{rgb}{0.5,0.5,0.5}
\definecolor{codepurple}{rgb}{0.58,0,0.82}
\definecolor{backcolour}{rgb}{0.95,0.95,0.92}

\lstdefinestyle{verilog}{
    language=Verilog, basicstyle=\footnotesize\ttfamily, numbers=left,
    numberstyle=\tiny\color{codegray}, stepnumber=1, numbersep=5pt,
    backgroundcolor=\color{backcolour}, frame=single, breaklines=true,
    tabsize=2, captionpos=b,
}
\lstdefinestyle{cstyle}{
    language=C, basicstyle=\footnotesize\ttfamily, numbers=left,
    numberstyle=\tiny\color{codegray}, stepnumber=1, numbersep=5pt,
    backgroundcolor=\color{backcolour}, frame=single, breaklines=true,
    tabsize=2, captionpos=b,
}
\lstdefinestyle{ascii}{
    basicstyle=\footnotesize\ttfamily, frame=single, breaklines=true,
}
\lstset{style=verilog}

\usepackage{bytefield}
\usepackage{tikz}
\usetikzlibrary{shapes,arrows,positioning,calc,fit,shadows}
\usepackage{pgfplots}
\pgfplotsset{compat=1.18}

\usepackage[backend=biber,style=ieee,sorting=none]{biblatex}
\addbibresource{references.bib}

\newcommand{\code}[1]{\texttt{#1}}
\newcommand{\reg}[1]{\texttt{#1}}

%------------------------------------------------------------------------------
\title{
    \vspace*{2cm}
    \Huge\textbf{RISC-V Custom Instruction Based \\[0.3cm]
    Lightweight CNN Accelerator \\[0.3cm]
    FPGA Prototype Validation}\\[2cm]
    \Large A Thesis Submitted in Partial Fulfillment \\[0.3cm]
    \Large of the Requirements for the Degree of \\[0.3cm]
    \Large Bachelor of Science in Computer Engineering \\[2cm]
    \large\textit{by}\\[1cm]
    \large\textbf{JU JIAXING}\\[0.3cm]
    \large Student ID: [Please fill in]\\[1.5cm]
    \large Supervisor: [Supervisor Name]\\[0.3cm]
    \large Department of Computer Science\\[0.3cm]
    \large City University of Hong Kong\\[2cm]
    \large\today
}

\begin{document}
\maketitle
\thispagestyle{empty}
\newpage

\chapter*{Declaration}
\addcontentsline{toc}{chapter}{Declaration}
\input{chapters/00_declaration.tex}
\newpage

\chapter*{Abstract}
\addcontentsline{toc}{chapter}{Abstract}
\input{chapters/00_abstract.tex}
\newpage

\tableofcontents
\newpage

\listoffigures
\addcontentsline{toc}{chapter}{List of Figures}
\listoftables
\addcontentsline{toc}{chapter}{List of Tables}
\newpage

\chapter{Introduction}
\input{chapters/01_introduction.tex}

\chapter{Background and Related Work}
\input{chapters/02_background.tex}

\chapter{System Architecture and Methodology}
\input{chapters/03_methodology.tex}

\chapter{Implementation and Results}
\input{chapters/04_results.tex}

\chapter{Discussion}
\input{chapters/05_discussion.tex}

\chapter{Conclusion and Future Work}
\input{chapters/06_conclusion.tex}

\printbibliography[heading=bibintoc, title=References]
\end{document}
"""
    with open(os.path.join(OUTPUT_DIR, 'main.tex'), 'w', encoding='utf-8') as f:
        f.write(content)
    print("  main.tex")

#-----------------------------------------------------------------------
def copy_figures():
    for fname in set(FIG_FILES.values()):
        src = os.path.join(FIGURES_SRC, fname)
        dst = os.path.join(FIGURES_DST, fname)
        if os.path.exists(src):
            shutil.copy2(src, dst)

def copy_bibliography():
    src = r"C:\Users\16084\Documents\Graduation_Design_Library\09_Thesis_Writing\References\references.bib"
    dst = os.path.join(OUTPUT_DIR, 'references.bib')
    if os.path.exists(src):
        shutil.copy2(src, dst)
        print(f"  copied references.bib ({os.path.getsize(dst)} bytes)")

def build_scripts():
    bat = r"""@echo off
echo === XeLaTeX (1) ===
xelatex -interaction=nonstopmode main.tex
echo === Biber ===
biber main
echo === XeLaTeX (2) ===
xelatex -interaction=nonstopmode main.tex
echo === XeLaTeX (3) ===
xelatex -interaction=nonstopmode main.tex
echo === Output: main.pdf ===
"""
    with open(os.path.join(OUTPUT_DIR, 'build.bat'), 'w') as f:
        f.write(bat)

    mk = "all: pdf\npdf:\n\txelatex main\n\tbiber main\n\txelatex main\n\txelatex main\nclean:\n\trm -f *.aux *.bbl *.bcf *.blg *.log *.out *.run.xml *.toc *.lof *.lot main.pdf chapters/*.aux\n"
    with open(os.path.join(OUTPUT_DIR, 'Makefile'), 'w') as f:
        f.write(mk)
    print("  build.bat, Makefile")

#-----------------------------------------------------------------------
if __name__ == '__main__':
    print("=== Generating LaTeX thesis ===\n")
    generate_main()

    print("Chapters:")
    generate_declaration()
    generate_abstract()
    chapters = [
        (1, 'Chapter 1. Introduction'), (2, 'Chapter 2. Background'),
        (3, 'Chapter 3. Methodology'), (4, 'Chapter 4. Results'),
        (5, 'Chapter 5. Discussion'), (6, 'Chapter 6. Conclusion'),
    ]
    for n, t in chapters:
        generate_chapter(n, t)

    print("Figures:")
    copy_figures()
    print("  all figures copied")

    print("Bibliography:")
    copy_bibliography()
    build_scripts()
    print("\n=== Done. Files in:", OUTPUT_DIR)
